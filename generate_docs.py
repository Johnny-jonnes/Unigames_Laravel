import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# CONFIGURATION GLOBALE DU STYLE
# ============================================================
PRIMARY_COLOR = RGBColor(26, 75, 173)
SECONDARY_COLOR = RGBColor(16, 185, 129)
DARK_COLOR = RGBColor(30, 41, 59)
MUTED_COLOR = RGBColor(100, 116, 139)
WHITE = RGBColor(255, 255, 255)

def setup_document(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_COLOR
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    colors_map = {
        'Heading 1': (Pt(22), PRIMARY_COLOR),
        'Heading 2': (Pt(16), RGBColor(40, 90, 190)),
        'Heading 3': (Pt(13), DARK_COLOR),
    }
    for heading, (size, color) in colors_map.items():
        if heading in doc.styles:
            h_style = doc.styles[heading]
            h_font = h_style.font
            h_font.name = 'Calibri'
            h_font.size = size
            h_font.color.rgb = color
            h_font.bold = True

def add_cover_page(doc, title, subtitle, version="1.0"):
    for _ in range(4):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title.upper())
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("_" * 60)
    run_line.font.color.rgb = SECONDARY_COLOR
    run_line.font.size = Pt(14)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(subtitle)
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = MUTED_COLOR
    run_sub.font.italic = True

    for _ in range(8):
        doc.add_paragraph()

    info_lines = [
        ("Projet", "UniGames - Plateforme de Gestion de Competitions Universitaires"),
        ("Client", "Direction des Sports Universitaires de Guinee"),
        ("Version", version),
        ("Date", "Juin 2026"),
        ("Auteur", "Equipe de Developpement UniGames"),
        ("Statut", "Document Final - Valide"),
    ]
    
    table = doc.add_table(rows=len(info_lines), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    for i, (label, value) in enumerate(info_lines):
        cell_l = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        run_l = cell_l.paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = MUTED_COLOR
        run_v = cell_v.paragraphs[0].add_run(value)
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = DARK_COLOR

    doc.add_page_break()

def add_toc_placeholder(doc):
    doc.add_heading("Table des Matieres", level=1)
    doc.add_paragraph("(A generer automatiquement via Word : References > Table des matieres)")
    doc.add_page_break()

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows_data):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, '1A4BAD')
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r, row in enumerate(rows_data):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            if r % 2 == 1:
                set_cell_shading(cell, 'F1F5F9')
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_COLOR

    doc.add_paragraph()

def add_bold_para(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    run_b = p.add_run(bold_text)
    run_b.font.bold = True
    run_b.font.color.rgb = DARK_COLOR
    run_n = p.add_run(normal_text)
    run_n.font.color.rgb = MUTED_COLOR

# ============================================================
# DOCUMENT 1 : CAHIER DES CHARGES (Enrichi)
# ============================================================
def create_cahier_des_charges():
    doc = Document()
    setup_document(doc)
    add_cover_page(doc, "Cahier des Charges", "Specifications et exigences du projet UniGames", "2.0")
    add_toc_placeholder(doc)

    # --- INTRODUCTION ---
    doc.add_heading("1. Introduction Generale", level=1)
    doc.add_heading("1.1. Contexte du Projet", level=2)
    doc.add_paragraph("Le sport universitaire en Guinee connait un essor considerable ces dernieres annees. Les competitions inter-universitaires, regroupant des dizaines de facultes et des centaines de joueurs a travers le pays, sont devenues des evenements majeurs de la vie etudiante. Cependant, l'organisation de ces tournois repose encore largement sur des processus manuels : tableaux Excel, feuilles de match papier, classements calcules a la main et communication informelle des resultats.")
    doc.add_paragraph("Ce mode de fonctionnement genere de nombreuses inefficacites : erreurs de calcul dans les classements, pertes de donnees, difficulte a obtenir des informations en temps reel, et absence d'historique fiable des competitions passees. Le projet UniGames nait de la volonte de moderniser et de digitaliser l'ensemble de ce processus.")

    doc.add_heading("1.2. Objectifs du Projet", level=2)
    doc.add_paragraph("Objectif principal : Concevoir et developper une plateforme web complete permettant la gestion centralisee des competitions sportives inter-universitaires guineennes.")
    doc.add_paragraph("Objectifs secondaires :", style='List Bullet')
    doc.add_paragraph("Offrir un tableau de bord en temps reel avec statistiques detaillees par edition", style='List Bullet')
    doc.add_paragraph("Permettre la gestion complete du cycle de vie d'un tournoi (creation, planification des matchs, saisie des scores, classements automatiques)", style='List Bullet')
    doc.add_paragraph("Assurer la tracabilite complete des resultats et des performances des joueurs", style='List Bullet')
    doc.add_paragraph("Proposer une interface utilisateur moderne, reactive et intuitive", style='List Bullet')
    doc.add_paragraph("Implementer un systeme de gestion des droits d'acces multi-roles (Administrateur, Staff, Visiteur)", style='List Bullet')

    doc.add_heading("1.3. Perimetre Fonctionnel", level=2)
    add_styled_table(doc,
        ["Module", "Description", "Priorite"],
        [
            ["Gestion des Editions", "Creation, modification et suivi des competitions annuelles. Gestion du statut (a venir, en cours, terminee).", "Haute"],
            ["Gestion des Facultes", "Enregistrement des etablissements participants avec logo et couleur identitaire.", "Haute"],
            ["Gestion des Disciplines", "Configuration des sports (Football, Basketball, Handball, Volleyball, Athletisme).", "Haute"],
            ["Gestion des Equipes", "Inscription des equipes par faculte, par discipline et par edition. Calcul automatique des statistiques.", "Haute"],
            ["Gestion des Joueurs", "Enregistrement des joueurs avec numero de maillot, sexe, et suivi des buts marques.", "Haute"],
            ["Gestion des Matchs", "Programmation des rencontres, saisie des scores et des buteurs, gestion des phases du tournoi.", "Critique"],
            ["Classements", "Calcul automatique des classements par discipline avec points, victoires, nuls, defaites et difference de buts.", "Haute"],
            ["Arbre du Tournoi", "Visualisation graphique de l'avancement des phases finales (Poules, Quarts, Demies, Finale).", "Moyenne"],
            ["Tableau de Bord", "Vue d'ensemble avec KPIs, derniers resultats, prochains matchs et meilleurs buteurs.", "Haute"],
            ["Gestion des Utilisateurs", "Administration des comptes staff avec attribution des roles et droits d'acces.", "Haute"],
            ["Authentification", "Systeme de connexion securise avec gestion de profil et reinitialisation de mot de passe.", "Critique"],
        ]
    )

    # --- ACTEURS ---
    doc.add_heading("2. Acteurs du Systeme", level=1)
    doc.add_heading("2.1. Matrice des Droits d'Acces", level=2)
    add_styled_table(doc,
        ["Fonctionnalite", "Administrateur", "Staff", "Visiteur (connecte)"],
        [
            ["Consulter le tableau de bord", "Oui", "Oui", "Oui"],
            ["Consulter les classements", "Oui", "Oui", "Oui"],
            ["Consulter l'arbre du tournoi", "Oui", "Oui", "Oui"],
            ["Voir les details d'un match", "Oui", "Oui", "Oui"],
            ["Creer/Modifier/Supprimer une Edition", "Oui", "Non", "Non"],
            ["Creer/Modifier/Supprimer une Faculte", "Oui", "Non", "Non"],
            ["Creer/Modifier/Supprimer une Discipline", "Oui", "Non", "Non"],
            ["Creer/Modifier/Supprimer une Equipe", "Oui", "Non", "Non"],
            ["Creer/Modifier un Joueur", "Oui", "Oui", "Non"],
            ["Supprimer un Joueur", "Oui", "Oui", "Non"],
            ["Programmer un Match", "Oui", "Oui", "Non"],
            ["Saisir le score d'un Match", "Oui", "Oui", "Non"],
            ["Supprimer un Match", "Oui", "Oui", "Non"],
            ["Gerer les comptes utilisateurs", "Oui", "Non", "Non"],
        ]
    )

    doc.add_heading("2.2. Description Detaillee des Roles", level=2)
    add_bold_para(doc, "Administrateur : ", "Dispose de l'acces complet a toutes les fonctionnalites de la plateforme. Il est responsable de la creation des editions, de l'ajout des facultes et disciplines, de la configuration des equipes et de la gestion des comptes utilisateurs. Il est le seul a pouvoir supprimer des entites majeures (editions, facultes, disciplines, equipes). Il peut egalement effectuer toutes les operations devolues au Staff.")
    add_bold_para(doc, "Staff : ", "Personnel de l'organisation du tournoi. Il peut creer et modifier des joueurs, programmer des matchs, saisir les scores et les buteurs. Il ne peut pas acceder aux fonctions d'administration systeme (gestion des utilisateurs) ni supprimer des entites structurelles.")
    add_bold_para(doc, "Visiteur connecte : ", "Tout utilisateur authentifie ne disposant pas de role specifique. Il beneficie d'un acces en lecture seule a l'ensemble des donnees publiques : tableau de bord, classements, calendrier des matchs, arbre du tournoi.")

    # --- EXIGENCES NON-FONCTIONNELLES ---
    doc.add_heading("3. Exigences Non-Fonctionnelles", level=1)
    doc.add_heading("3.1. Performance", level=2)
    doc.add_paragraph("Le temps de reponse moyen pour l'affichage d'une page ne doit pas exceder 2 secondes en conditions normales d'utilisation.", style='List Bullet')
    doc.add_paragraph("La plateforme doit supporter au minimum 50 utilisateurs simultanement connectes.", style='List Bullet')
    doc.add_paragraph("Les calculs de classement doivent etre effectues en temps reel, sans cache, pour garantir la fraicheur des donnees.", style='List Bullet')

    doc.add_heading("3.2. Securite", level=2)
    doc.add_paragraph("Authentification par email et mot de passe avec hashage bcrypt.", style='List Bullet')
    doc.add_paragraph("Protection CSRF sur tous les formulaires (token Laravel natif).", style='List Bullet')
    doc.add_paragraph("Middleware de controle d'acces par role sur chaque groupe de routes.", style='List Bullet')
    doc.add_paragraph("Validation cote serveur de toutes les entrees utilisateur.", style='List Bullet')

    doc.add_heading("3.3. Ergonomie et Accessibilite", level=2)
    doc.add_paragraph("Interface responsive adaptee aux ecrans desktop, tablette et mobile.", style='List Bullet')
    doc.add_paragraph("Design moderne utilisant un systeme de design unifie (enterprise-card, enterprise-btn, enterprise-input).", style='List Bullet')
    doc.add_paragraph("Recherche en temps reel sur toutes les listes (Alpine.js).", style='List Bullet')
    doc.add_paragraph("Retours visuels immediats (messages de succes, confirmation avant suppression).", style='List Bullet')

    doc.add_heading("3.4. Maintenabilite", level=2)
    doc.add_paragraph("Architecture MVC stricte (Laravel). Separation claire des responsabilites.", style='List Bullet')
    doc.add_paragraph("Code source versionne avec Git, heberge sur GitHub.", style='List Bullet')
    doc.add_paragraph("Seeders fournis pour l'initialisation rapide d'un environnement de test avec donnees realistes.", style='List Bullet')

    # --- CONTRAINTES ---
    doc.add_heading("4. Contraintes Techniques", level=1)
    add_styled_table(doc,
        ["Composant", "Technologie", "Version"],
        [
            ["Langage Backend", "PHP", "8.4+"],
            ["Framework Backend", "Laravel", "13.7"],
            ["Base de Donnees", "MySQL", "8.0+"],
            ["Frontend - Templates", "Blade (Laravel)", "-"],
            ["Frontend - Reactivite", "Alpine.js", "3.x"],
            ["Frontend - Style", "Tailwind CSS", "3.x"],
            ["Build Frontend", "Vite", "8.x"],
            ["Authentification", "Laravel Breeze", "-"],
            ["Serveur Web (dev)", "artisan serve", "-"],
            ["Gestionnaire de paquets PHP", "Composer", "2.x"],
            ["Gestionnaire de paquets JS", "NPM", "18+"],
        ]
    )

    doc.save("docs/1_Cahier_des_Charges.docx")
    print("[OK] 1_Cahier_des_Charges.docx")

# ============================================================
# DOCUMENT 2 : MANUEL UTILISATEUR (Hyper-detaille)
# ============================================================
def create_manuel_utilisateur():
    doc = Document()
    setup_document(doc)
    add_cover_page(doc, "Manuel Utilisateur", "Guide Complet d'Utilisation de la Plateforme UniGames", "2.0")
    add_toc_placeholder(doc)

    # --- PRISE EN MAIN ---
    doc.add_heading("1. Prise en Main de la Plateforme", level=1)
    doc.add_heading("1.1. Acces a l'Application", level=2)
    doc.add_paragraph("UniGames est une application web accessible depuis n'importe quel navigateur moderne (Chrome, Firefox, Edge, Safari). Aucune installation locale n'est necessaire pour les utilisateurs finaux.")
    doc.add_paragraph("URL d'acces : http://localhost:8001 (environnement de developpement)")
    doc.add_paragraph("A l'ouverture, l'utilisateur est automatiquement redirige vers la page de connexion s'il n'est pas deja authentifie.")

    doc.add_heading("1.2. Page de Connexion", level=2)
    doc.add_paragraph("La page de connexion presente un formulaire epure avec deux champs :")
    doc.add_paragraph("Adresse e-mail : Saisissez l'adresse email associee a votre compte (ex: admin@unigames.com).", style='List Bullet')
    doc.add_paragraph("Mot de passe : Saisissez votre mot de passe. Un lien 'Mot de passe oublie ?' est disponible pour la reinitialisation.", style='List Bullet')
    doc.add_paragraph("Apres authentification reussie, vous etes redirige vers le Tableau de Bord.")
    doc.add_paragraph("En cas d'erreur de saisie, un message d'erreur en rouge apparait sous le champ concerne.")
    if os.path.exists("docs/images/2_login.jpeg"):
        doc.add_picture("docs/images/2_login.jpeg", width=Inches(6.0))

    doc.add_heading("1.3. Navigation Principale", level=2)
    doc.add_paragraph("La barre laterale gauche (sidebar) constitue le point d'entree principal vers tous les modules de la plateforme. Elle est visible en permanence et contient les liens suivants :")
    add_styled_table(doc,
        ["Icone", "Menu", "Description", "Acces"],
        [
            ["Grille", "Tableau de Bord", "Vue d'ensemble avec statistiques, derniers resultats et prochains matchs", "Tous"],
            ["Trophee", "Editions", "Liste des competitions. Cliquer pour voir les details d'une edition.", "Tous"],
            ["Batiment", "Facultes", "Liste des etablissements participants avec leurs equipes.", "Tous"],
            ["Ballon", "Disciplines", "Sports disponibles dans la competition.", "Tous"],
            ["Equipe", "Equipes", "Toutes les equipes inscrites, filtrees par edition.", "Tous"],
            ["Personne", "Joueurs", "Liste de tous les joueurs avec leur equipe, faculte et statistiques.", "Tous"],
            ["Calendrier", "Matchs", "Calendrier complet des rencontres, groupees par discipline.", "Tous"],
            ["Classement", "Classements", "Tableaux de classement par discipline et meilleurs buteurs.", "Tous"],
            ["Engrenage", "Gestion Staff", "Creation et gestion des comptes utilisateurs.", "Admin uniquement"],
        ]
    )

    # --- TABLEAU DE BORD ---
    doc.add_heading("2. Module : Tableau de Bord", level=1)
    doc.add_heading("2.1. Presentation Generale", level=2)
    doc.add_paragraph("Le tableau de bord est la premiere page affichee apres la connexion. Il offre une vue synthetique de l'edition selectionnee. En haut a droite, un selecteur d'edition permet de basculer entre les differentes competitions (ex : 'Edition 2026').")
    if os.path.exists("docs/images/4_dashboard_principal.jpeg"):
        doc.add_picture("docs/images/4_dashboard_principal.jpeg", width=Inches(6.0))
    
    doc.add_heading("2.2. Indicateurs Cles (KPIs)", level=2)
    doc.add_paragraph("La partie superieure affiche des cartes de statistiques avec les indicateurs suivants :")
    add_styled_table(doc,
        ["KPI", "Description", "Exemple"],
        [
            ["Facultes", "Nombre d'etablissements inscrits pour cette edition", "8 facultes"],
            ["Equipes", "Nombre total d'equipes inscrites (toutes disciplines confondues)", "24 equipes"],
            ["Matchs Joues", "Nombre de matchs deja termines sur le total programme", "30/42 matchs"],
            ["Joueurs", "Nombre total de joueurs enregistres pour cette edition", "192 joueurs"],
        ]
    )
    
    doc.add_heading("2.3. Derniers Resultats", level=2)
    doc.add_paragraph("Un tableau affiche les 5 derniers matchs joues avec le score final, les equipes concernees et la discipline. Chaque ligne est cliquable et redirige vers la fiche detaillee du match.")
    
    doc.add_heading("2.4. Prochains Matchs", level=2)
    doc.add_paragraph("Un second tableau montre les 5 prochains matchs planifies avec la date, l'heure, le lieu et les equipes. Si l'edition est terminee, cette section est vide.")
    
    doc.add_heading("2.5. Meilleurs Buteurs", level=2)
    doc.add_paragraph("Un classement des 5 meilleurs buteurs de l'edition selectionnee, affichant le nom du joueur, son equipe, sa faculte et le nombre de buts marques.")

    # --- GESTION DES EDITIONS ---
    doc.add_heading("3. Module : Gestion des Editions", level=1)
    doc.add_heading("3.1. Liste des Editions", level=2)
    doc.add_paragraph("La page principale affiche toutes les editions existantes sous forme de cartes visuelles. Chaque carte contient le nom de l'edition, ses dates de debut et de fin, son statut (badge colore) et le nombre d'equipes et de matchs associes.")
    doc.add_paragraph("Statuts possibles :")
    doc.add_paragraph("A venir (badge bleu) : L'edition n'a pas encore commence.", style='List Bullet')
    doc.add_paragraph("En cours (badge jaune/orange) : L'edition est active, les matchs sont en cours de deroulement.", style='List Bullet')
    doc.add_paragraph("Terminee (badge vert) : Tous les matchs ont ete joues. L'edition est archivee.", style='List Bullet')
    if os.path.exists("docs/images/3_selection_edition.jpeg"):
        doc.add_picture("docs/images/3_selection_edition.jpeg", width=Inches(6.0))

    doc.add_heading("3.2. Creer une Edition (Admin)", level=2)
    doc.add_paragraph("Cliquer sur le bouton '+ Nouvelle Edition' dans l'en-tete de la page.")
    doc.add_paragraph("Remplir le formulaire avec les champs suivants :")
    add_styled_table(doc,
        ["Champ", "Type", "Obligatoire", "Description"],
        [
            ["Nom", "Texte", "Oui", "Nom de l'edition (ex: 'Edition 2026')"],
            ["Date de debut", "Date", "Oui", "Date d'ouverture du tournoi"],
            ["Date de fin", "Date", "Oui", "Date de cloture du tournoi"],
            ["Lieu", "Texte", "Non", "Lieu principal du tournoi (ex: 'Conakry')"],
            ["Description", "Texte long", "Non", "Description libre de l'edition"],
            ["Statut", "Selection", "Oui", "a_venir, en_cours ou terminee"],
        ]
    )

    doc.add_heading("3.3. Page de Detail d'une Edition", level=2)
    doc.add_paragraph("En cliquant sur une edition, on accede a sa page de detail qui affiche :")
    doc.add_paragraph("Les informations generales (nom, dates, lieu, description, statut)", style='List Bullet')
    doc.add_paragraph("Le nombre de facultes inscrites, d'equipes participantes et de matchs programmes", style='List Bullet')
    doc.add_paragraph("Un bouton 'Arbre du Tournoi' permettant de visualiser graphiquement les phases finales", style='List Bullet')
    doc.add_paragraph("Un bouton 'Modifier' (admin uniquement) pour editer les informations de l'edition", style='List Bullet')

    # --- GESTION DES FACULTES ---
    doc.add_heading("4. Module : Gestion des Facultes", level=1)
    doc.add_heading("4.1. Liste des Facultes", level=2)
    doc.add_paragraph("Cette page affiche l'ensemble des etablissements universitaires inscrits a l'edition selectionnee. Un champ de recherche en temps reel permet de filtrer les resultats par nom. Chaque carte de faculte affiche le nom, le logo (initiales stylisees), la couleur identitaire et le nombre d'equipes inscrites.")
    if os.path.exists("docs/images/5_liste_facultes.jpeg"):
        doc.add_picture("docs/images/5_liste_facultes.jpeg", width=Inches(6.0))
    
    doc.add_heading("4.2. Creer une Faculte (Admin)", level=2)
    doc.add_paragraph("Cliquer sur '+ Nouvelle Faculte'. Remplir les champs suivants :")
    add_styled_table(doc,
        ["Champ", "Type", "Obligatoire", "Description"],
        [
            ["Nom", "Texte", "Oui", "Nom complet (ex: 'Universite Gamal Abdel Nasser de Conakry')"],
            ["Logo", "Image", "Non", "Upload du logo de l'etablissement"],
            ["Couleur", "Selecteur couleur", "Non", "Couleur identitaire pour les badges et graphiques"],
        ]
    )
    if os.path.exists("docs/images/6_ajouter_faculte.jpeg"):
        doc.add_picture("docs/images/6_ajouter_faculte.jpeg", width=Inches(6.0))

    doc.add_heading("4.3. Detail d'une Faculte", level=2)
    doc.add_paragraph("La page de detail affiche les informations de la faculte ainsi que la liste de toutes ses equipes inscrites, avec pour chaque equipe : la discipline, le nombre de joueurs et un lien vers la fiche de l'equipe.")
    if os.path.exists("docs/images/7_details_faculte.jpeg"):
        doc.add_picture("docs/images/7_details_faculte.jpeg", width=Inches(6.0))

    # --- GESTION DES EQUIPES ---
    doc.add_heading("5. Module : Gestion des Equipes", level=1)
    doc.add_heading("5.1. Liste des Equipes", level=2)
    doc.add_paragraph("La page affiche toutes les equipes de l'edition en cours sous forme de tableau. Chaque ligne indique le nom de l'equipe, sa faculte d'appartenance, la discipline pratiquee et le nombre de joueurs inscrits. Une barre de recherche filtre les resultats en temps reel.")
    if os.path.exists("docs/images/11_repertoire_equipes.jpeg"):
        doc.add_picture("docs/images/11_repertoire_equipes.jpeg", width=Inches(6.0))

    doc.add_heading("5.2. Inscription d'une Equipe (Admin)", level=2)
    doc.add_paragraph("Le formulaire de creation d'equipe requiert :")
    add_styled_table(doc,
        ["Champ", "Type", "Obligatoire", "Description"],
        [
            ["Nom", "Texte", "Oui", "Nom de l'equipe (souvent genere automatiquement a partir de la faculte et de la discipline)"],
            ["Faculte", "Selection", "Oui", "L'etablissement representant cette equipe"],
            ["Discipline", "Selection", "Oui", "Le sport pratique par cette equipe"],
            ["Edition", "Selection", "Oui", "L'edition de la competition"],
        ]
    )
    doc.add_paragraph("Regle de gestion : Une faculte ne peut inscrire qu'une seule equipe par discipline et par edition. En cas de doublon, le systeme refuse la creation avec un message d'erreur.")
    if os.path.exists("docs/images/9_ajouter_equipe.jpeg"):
        doc.add_picture("docs/images/9_ajouter_equipe.jpeg", width=Inches(6.0))

    doc.add_heading("5.3. Fiche Detaillee d'une Equipe", level=2)
    doc.add_paragraph("La page de detail d'une equipe affiche :")
    doc.add_paragraph("Les informations generales (nom, faculte, discipline, edition)", style='List Bullet')
    doc.add_paragraph("Les statistiques de performance : matchs joues, victoires, nuls, defaites, buts marques, buts encaisses, difference de buts, points", style='List Bullet')
    doc.add_paragraph("La liste complete des joueurs de l'equipe avec leur numero de maillot et le nombre de buts marques", style='List Bullet')
    doc.add_paragraph("L'historique des matchs joues par cette equipe avec les scores", style='List Bullet')

    # --- GESTION DES JOUEURS ---
    doc.add_heading("6. Module : Gestion des Joueurs", level=1)
    doc.add_heading("6.1. Liste des Joueurs", level=2)
    doc.add_paragraph("La page des joueurs affiche un tableau complet de tous les sportifs inscrits pour l'edition selectionnee. Les colonnes affichees sont : Nom, Prenom, Numero de Maillot, Sexe, Equipe, Faculte, et Nombre de Buts. Un champ de recherche en haut permet de filtrer instantanement par nom ou equipe.")
    if os.path.exists("docs/images/12_repertoire_joueurs.jpeg"):
        doc.add_picture("docs/images/12_repertoire_joueurs.jpeg", width=Inches(6.0))

    doc.add_heading("6.2. Enregistrer un Joueur (Admin/Staff)", level=2)
    add_styled_table(doc,
        ["Champ", "Type", "Obligatoire", "Validation"],
        [
            ["Nom", "Texte", "Oui", "Max 255 caracteres"],
            ["Prenom", "Texte", "Oui", "Max 255 caracteres"],
            ["Sexe", "Selection (M/F)", "Oui", "Valeurs autorisees : M ou F"],
            ["Equipe", "Selection", "Oui", "Doit exister dans la table equipes"],
            ["Numero de Maillot", "Entier", "Non", "Minimum 1"],
        ]
    )
    if os.path.exists("docs/images/10_ajouter_joueur.jpeg"):
        doc.add_picture("docs/images/10_ajouter_joueur.jpeg", width=Inches(6.0))

    doc.add_heading("6.3. Actions Disponibles", level=2)
    doc.add_paragraph("Modifier : Le bouton bleu 'Modifier' ouvre le formulaire de modification pre-rempli avec les donnees actuelles du joueur.", style='List Bullet')
    doc.add_paragraph("Supprimer : Le bouton rouge 'Supprimer' declenche une boite de dialogue de confirmation. La suppression est irreversible.", style='List Bullet')
    doc.add_paragraph("Voir la fiche : Cliquer sur la ligne du joueur pour acceder a sa fiche detaillee.", style='List Bullet')

    # --- GESTION DES MATCHS ---
    doc.add_heading("7. Module : Gestion des Matchs", level=1)
    doc.add_heading("7.1. Calendrier des Rencontres", level=2)
    doc.add_paragraph("La page des matchs est organisee par discipline sportive. Chaque discipline est representee par un panneau repliable (accordeon) qui, une fois ouvert, affiche la liste de tous les matchs de cette discipline pour l'edition en cours. Les matchs sont tries par date (du plus recent au plus ancien).")
    doc.add_paragraph("Pour chaque match, le tableau affiche :")
    doc.add_paragraph("La date et l'heure du match, ainsi que le lieu", style='List Bullet')
    doc.add_paragraph("L'equipe A (domicile) et l'equipe B (visiteur) avec le nom de leur faculte", style='List Bullet')
    doc.add_paragraph("Le score (si le match est joue) ou 'VS' (si planifie)", style='List Bullet')
    doc.add_paragraph("Le statut du match : badge 'Joue' (vert), 'En cours' (orange) ou 'Planifie' (bleu)", style='List Bullet')
    doc.add_paragraph("Les boutons d'action : Voir, Modifier (bleu), Supprimer (rouge)", style='List Bullet')
    if os.path.exists("docs/images/13_calendrier_rencontres.jpeg"):
        doc.add_picture("docs/images/13_calendrier_rencontres.jpeg", width=Inches(6.0))

    doc.add_heading("7.2. Programmer un Match (Admin/Staff)", level=2)
    doc.add_paragraph("Cliquer sur le bouton '+ Programmer un match' dans l'en-tete de la page. Ce bouton n'apparait que si l'edition selectionnee n'est pas terminee.")
    add_styled_table(doc,
        ["Champ", "Type", "Obligatoire", "Description"],
        [
            ["Edition", "Selection", "Oui", "L'edition du tournoi"],
            ["Discipline", "Selection", "Oui", "Le sport concerne"],
            ["Equipe A", "Selection", "Oui", "L'equipe recevante. Doit etre differente de l'Equipe B."],
            ["Equipe B", "Selection", "Oui", "L'equipe visiteuse"],
            ["Date et Heure", "Datetime", "Oui", "Moment de la rencontre"],
            ["Lieu", "Texte", "Non", "Stade ou terrain (ex: 'Stade du 28 Septembre')"],
            ["Phase", "Selection", "Oui", "Phase du tournoi : Poules, Quarts de Finale, Demi-Finales, Petite Finale ou Finale"],
        ]
    )
    doc.add_paragraph("Validation : Le systeme refuse de programmer un match pour une edition au statut 'terminee'. Un message d'erreur est affiche.")

    doc.add_heading("7.3. Saisir le Score d'un Match", level=2)
    doc.add_paragraph("Cette fonctionnalite est accessible depuis la fiche detaillee d'un match planifie. Le formulaire de saisie du score est divise en deux colonnes (Equipe A et Equipe B) et comprend :")
    doc.add_paragraph("Un champ numerique pour le score de chaque equipe (affichage en grand format, style monospace).", style='List Bullet')
    doc.add_paragraph("Une section 'Buteurs' pour chaque equipe : permet d'ajouter dynamiquement des buteurs en selectionnant le joueur dans une liste deroulante et en indiquant le nombre de buts marques.", style='List Bullet')
    doc.add_paragraph("Le bouton '+ Ajouter un buteur' permet d'ajouter autant de lignes que necessaire.", style='List Bullet')
    doc.add_paragraph("La croix rouge a cote de chaque buteur permet de retirer un buteur ajoute par erreur.", style='List Bullet')
    doc.add_paragraph("Lorsque le formulaire est valide et soumis, le systeme effectue automatiquement les operations suivantes :")
    doc.add_paragraph("1. Le statut du match passe a 'joue'.", style='List Number')
    doc.add_paragraph("2. Les scores sont enregistres dans la table matchs.", style='List Number')
    doc.add_paragraph("3. Les buteurs sont stockes au format JSON dans le champ 'buteurs'.", style='List Number')
    doc.add_paragraph("4. Le compteur de buts de chaque joueur concerne est automatiquement incremente dans la table joueurs.", style='List Number')
    doc.add_paragraph("5. Les classements sont recalcules en temps reel lors de la prochaine consultation.", style='List Number')

    doc.add_heading("7.4. Fiche Detaillee d'un Match", level=2)
    doc.add_paragraph("La page de detail d'un match joue presente un 'scoreboard' central avec les logos des deux equipes, le score en grand format, la liste des buteurs de chaque cote, ainsi que les informations contextuelles (date, lieu, phase, discipline). Un badge 'Termine' confirme visuellement que le match est clos.")
    if os.path.exists("docs/images/14_details_match.jpeg"):
        doc.add_picture("docs/images/14_details_match.jpeg", width=Inches(6.0))

    # --- CLASSEMENTS ---
    doc.add_heading("8. Module : Classements", level=1)
    doc.add_heading("8.1. Classement par Discipline", level=2)
    doc.add_paragraph("La page de classement permet de selectionner une discipline sportive via un menu deroulant. Une fois la discipline choisie, un tableau de classement complet s'affiche avec les colonnes suivantes :")
    add_styled_table(doc,
        ["Colonne", "Abreviation", "Description"],
        [
            ["Position", "#", "Rang de l'equipe dans le classement"],
            ["Equipe", "-", "Nom de l'equipe et faculte"],
            ["Matchs Joues", "MJ", "Nombre total de matchs joues"],
            ["Victoires", "V", "Nombre de matchs gagnes"],
            ["Nuls", "N", "Nombre de matchs nuls"],
            ["Defaites", "D", "Nombre de matchs perdus"],
            ["Buts Marques", "BM", "Total des buts marques"],
            ["Buts Encaisses", "BE", "Total des buts encaisses"],
            ["Difference de Buts", "DB", "BM - BE"],
            ["Points", "Pts", "V*3 + N*1 + D*0"],
        ]
    )
    doc.add_paragraph("Le classement est trie automatiquement par nombre de points decroissant, puis par difference de buts en cas d'egalite.")
    if os.path.exists("docs/images/15_classements_statistiques.jpeg"):
        doc.add_picture("docs/images/15_classements_statistiques.jpeg", width=Inches(6.0))

    doc.add_heading("8.2. Tableau des Meilleurs Buteurs", level=2)
    doc.add_paragraph("En bas de la page de classement, un tableau affiche le top 10 des meilleurs buteurs de l'edition, toutes disciplines confondues. Pour chaque joueur : nom, prenom, equipe, faculte, discipline et nombre de buts.")

    # --- ARBRE DU TOURNOI ---
    doc.add_heading("9. Module : Arbre du Tournoi", level=1)
    doc.add_paragraph("L'arbre du tournoi est accessible depuis la page de detail d'une edition. Il offre une visualisation horizontale et interactive de l'avancement des phases finales pour chaque discipline :")
    doc.add_paragraph("Phase de Poules : Cartes compactes affichant les resultats de chaque match de poule.", style='List Bullet')
    doc.add_paragraph("Quarts de Finale : Cartes plus larges avec score et statut.", style='List Bullet')
    doc.add_paragraph("Demi-Finales : Cartes avec fond colore pour le vainqueur.", style='List Bullet')
    doc.add_paragraph("Petite Finale (3eme place) : Carte avec bordure en pointilles.", style='List Bullet')
    doc.add_paragraph("Grande Finale : Carte doree premium avec couronne emoji pour le champion. Un label 'OR' et un fond degrade ambre signalent l'importance de ce match.", style='List Bullet')
    doc.add_paragraph("Chaque carte de match est cliquable et redirige vers la fiche detaillee du match correspondant.")
    if os.path.exists("docs/images/18_arbre_tournoi.jpeg"):
        doc.add_picture("docs/images/18_arbre_tournoi.jpeg", width=Inches(6.0))

    # --- GESTION DES UTILISATEURS ---
    doc.add_heading("10. Module : Gestion des Utilisateurs (Admin)", level=1)
    doc.add_heading("10.1. Liste des Utilisateurs", level=2)
    doc.add_paragraph("Accessible uniquement aux administrateurs via le menu 'Gestion Staff'. Cette page affiche un tableau de tous les comptes utilisateurs avec : nom, email, role (badge colore) et date d'inscription. Des boutons 'Modifier' et 'Supprimer' sont disponibles pour chaque compte.")
    if os.path.exists("docs/images/16_gestion_staff.jpeg"):
        doc.add_picture("docs/images/16_gestion_staff.jpeg", width=Inches(6.0))

    doc.add_heading("10.2. Creer un Compte Staff", level=2)
    add_styled_table(doc,
        ["Champ", "Type", "Obligatoire", "Description"],
        [
            ["Nom", "Texte", "Oui", "Nom complet de l'utilisateur"],
            ["Email", "Email", "Oui", "Adresse email unique (servira d'identifiant de connexion)"],
            ["Mot de passe", "Mot de passe", "Oui", "Minimum 8 caracteres, sera hashe en bcrypt"],
            ["Role", "Selection", "Oui", "admin ou staff"],
        ]
    )

    doc.add_heading("10.3. Modifier un Compte", level=2)
    doc.add_paragraph("L'administrateur peut modifier le nom, l'email et le role de n'importe quel utilisateur. Le mot de passe peut etre reinitialise en saisissant un nouveau mot de passe dans le champ dedie (laisser vide pour ne pas modifier).")

    # --- PROFIL ---
    doc.add_heading("11. Gestion du Profil Personnel", level=1)
    doc.add_paragraph("Chaque utilisateur connecte peut acceder a son profil via l'icone utilisateur en haut a droite. La page de profil offre trois sections :")
    doc.add_paragraph("Informations du Profil : Modification du nom et de l'email.", style='List Bullet')
    doc.add_paragraph("Modification du Mot de Passe : Saisie de l'ancien mot de passe, puis du nouveau mot de passe avec confirmation.", style='List Bullet')
    doc.add_paragraph("Suppression du Compte : Action irreversible necessitant la saisie du mot de passe actuel pour confirmation.", style='List Bullet')
    if os.path.exists("docs/images/17_profil_parametres.jpeg"):
        doc.add_picture("docs/images/17_profil_parametres.jpeg", width=Inches(6.0))

    doc.save("docs/2_Manuel_Utilisateur.docx")
    print("[OK] 2_Manuel_Utilisateur.docx")

# ============================================================
# DOCUMENT 3 : DOSSIER D'ARCHITECTURE
# ============================================================
def create_dossier_architecture():
    doc = Document()
    setup_document(doc)
    add_cover_page(doc, "Dossier d'Architecture", "Architecture Technique, Base de Donnees et Couches Applicatives", "2.0")
    add_toc_placeholder(doc)

    doc.add_heading("1. Vue d'Ensemble de l'Architecture", level=1)
    doc.add_heading("1.1. Patron Architectural : MVC (Model-View-Controller)", level=2)
    doc.add_paragraph("UniGames adopte le patron MVC impose par le framework Laravel. Ce patron separe strictement les responsabilites en trois couches :")
    add_bold_para(doc, "Modele (Model) : ", "Represente les entites metier et encapsule la logique d'acces aux donnees. Chaque modele Eloquent correspond a une table de la base de donnees et definit les relations, les accesseurs calcules (ex: points, victoires) et les regles de validation.")
    add_bold_para(doc, "Vue (View) : ", "Les templates Blade (.blade.php) generent le HTML envoye au navigateur. Ils utilisent le systeme de composants de Laravel (x-app-layout, x-sidebar, x-topbar) pour garantir la coherence visuelle sur toutes les pages.")
    add_bold_para(doc, "Controleur (Controller) : ", "Les controleurs orchestrent le flux applicatif : reception des requetes HTTP, validation des donnees, appel aux modeles, preparation des donnees et retour de la vue appropriee.")

    doc.add_heading("1.2. Diagramme d'Architecture en Couches", level=2)
    doc.add_paragraph("[Navigateur] <-> [Vite/HMR] <-> [Blade + Alpine.js + Tailwind CSS]")
    doc.add_paragraph("                                      |")
    doc.add_paragraph("                               [Routes (web.php)]")
    doc.add_paragraph("                                      |")
    doc.add_paragraph("                         [Middlewares (auth, admin, can.manage)]")
    doc.add_paragraph("                                      |")
    doc.add_paragraph("                             [Controllers Laravel]")
    doc.add_paragraph("                                      |")
    doc.add_paragraph("                          [Models Eloquent + Relations]")
    doc.add_paragraph("                                      |")
    doc.add_paragraph("                               [MySQL Database]")

    # --- MODELES ---
    doc.add_heading("2. Couche Modele (Eloquent ORM)", level=1)
    doc.add_heading("2.1. Inventaire des Modeles", level=2)
    add_styled_table(doc,
        ["Modele", "Table", "Attributs Fillable", "Relations"],
        [
            ["User", "users", "name, email, password, role", "-"],
            ["Edition", "editions", "nom, date_debut, date_fin, lieu, description, statut", "hasMany: Equipe, Match_, Faculte"],
            ["Faculte", "facultes", "nom, logo, couleur, edition_id", "belongsTo: Edition | hasMany: Equipe"],
            ["Discipline", "disciplines", "nom, type, nb_joueurs", "hasMany: Equipe, Match_"],
            ["Equipe", "equipes", "nom, faculte_id, discipline_id, edition_id", "belongsTo: Faculte, Discipline, Edition | hasMany: Joueur, Match_"],
            ["Joueur", "joueurs", "nom, prenom, sexe, equipe_id, numero_maillot, buts", "belongsTo: Equipe"],
            ["Match_", "matchs", "equipe_a_id, equipe_b_id, discipline_id, edition_id, date_match, lieu, phase, score_a, score_b, statut, buteurs", "belongsTo: Equipe (x2), Discipline, Edition"],
        ]
    )

    doc.add_heading("2.2. Accesseurs Calcules (Modele Equipe)", level=2)
    doc.add_paragraph("Le modele Equipe implemente des accesseurs Eloquent qui calculent dynamiquement les statistiques de performance a partir des matchs joues :")
    add_styled_table(doc,
        ["Accesseur", "Formule", "Description"],
        [
            ["getPointsAttribute()", "V*3 + N*1", "Total des points cumules"],
            ["getMatchsJouesAttribute()", "COUNT(matchs WHERE statut=joue)", "Nombre de matchs termines"],
            ["getVictoiresAttribute()", "COUNT(matchs WHERE score_equipe > score_adverse)", "Nombre de victoires"],
            ["getNulsAttribute()", "COUNT(matchs WHERE score_a = score_b)", "Nombre de matchs nuls"],
            ["getDefaitesAttribute()", "COUNT(matchs WHERE score_equipe < score_adverse)", "Nombre de defaites"],
            ["getButsMarquesAttribute()", "SUM(scores de l'equipe)", "Total des buts inscrits"],
            ["getButsEncaissesAttribute()", "SUM(scores adverses)", "Total des buts encaisses"],
            ["getDifferenceButs()", "buts_marques - buts_encaisses", "Difference de buts"],
        ]
    )

    doc.add_heading("2.3. Particularite : Modele Match_", level=2)
    doc.add_paragraph("Le mot 'match' etant reserve en PHP, le modele est nomme Match_ (avec underscore). La propriete $table est explicitement definie a 'matchs' pour eviter toute ambiguite.")
    doc.add_paragraph("Le champ 'buteurs' utilise un cast JSON (array) pour stocker la structure suivante :")
    doc.add_paragraph('{ "equipe_a": [{"id": 45, "nb_buts": 2}, {"id": 48, "nb_buts": 1}], "equipe_b": [{"id": 62, "nb_buts": 1}] }')

    # --- BASE DE DONNEES ---
    doc.add_heading("3. Base de Donnees (MySQL)", level=1)
    doc.add_heading("3.1. Schema Relationnel Complet", level=2)
    doc.add_paragraph("Le diagramme ER complet est disponible dans le fichier 'docs/diagramme_architecture.html' (interactif, ouvrir dans un navigateur).")
    
    doc.add_heading("3.2. Dictionnaire de Donnees", level=2)
    doc.add_paragraph("Table : editions")
    add_styled_table(doc,
        ["Colonne", "Type", "Null", "Cle", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK, Auto-increment", "Identifiant unique"],
            ["nom", "VARCHAR(255)", "Non", "-", "Nom de l'edition"],
            ["date_debut", "DATE", "Non", "-", "Date de debut"],
            ["date_fin", "DATE", "Non", "-", "Date de fin"],
            ["lieu", "VARCHAR(255)", "Oui", "-", "Lieu principal"],
            ["description", "TEXT", "Oui", "-", "Description libre"],
            ["statut", "VARCHAR(50)", "Non", "-", "a_venir / en_cours / terminee"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    doc.add_paragraph("Table : matchs")
    add_styled_table(doc,
        ["Colonne", "Type", "Null", "Cle", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK", "Identifiant unique"],
            ["equipe_a_id", "BIGINT UNSIGNED", "Non", "FK -> equipes.id", "Equipe recevante"],
            ["equipe_b_id", "BIGINT UNSIGNED", "Non", "FK -> equipes.id", "Equipe visiteuse"],
            ["discipline_id", "BIGINT UNSIGNED", "Non", "FK -> disciplines.id", "Discipline du match"],
            ["edition_id", "BIGINT UNSIGNED", "Non", "FK -> editions.id", "Edition du tournoi"],
            ["date_match", "DATETIME", "Non", "-", "Date et heure du match"],
            ["lieu", "VARCHAR(255)", "Oui", "-", "Stade ou terrain"],
            ["phase", "VARCHAR(255)", "Non", "-", "Phase du tournoi"],
            ["score_a", "INT", "Oui", "-", "Score equipe A"],
            ["score_b", "INT", "Oui", "-", "Score equipe B"],
            ["statut", "VARCHAR(50)", "Non", "-", "planifie / en_cours / joue"],
            ["buteurs", "JSON", "Oui", "-", "Donnees JSON des buteurs"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    doc.add_paragraph("Table : joueurs")
    add_styled_table(doc,
        ["Colonne", "Type", "Null", "Cle", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK", "Identifiant unique"],
            ["nom", "VARCHAR(255)", "Non", "-", "Nom de famille"],
            ["prenom", "VARCHAR(255)", "Non", "-", "Prenom"],
            ["sexe", "CHAR(1)", "Non", "-", "M ou F"],
            ["equipe_id", "BIGINT UNSIGNED", "Non", "FK -> equipes.id", "Equipe d'appartenance"],
            ["numero_maillot", "INT", "Oui", "-", "Numero du maillot"],
            ["buts", "INT", "Non (defaut 0)", "-", "Total de buts marques"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    doc.add_paragraph("Table : equipes")
    add_styled_table(doc,
        ["Colonne", "Type", "Null", "Cle", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK", "Identifiant unique"],
            ["nom", "VARCHAR(255)", "Non", "-", "Nom de l'equipe"],
            ["faculte_id", "BIGINT UNSIGNED", "Non", "FK -> facultes.id", "Faculte representee"],
            ["discipline_id", "BIGINT UNSIGNED", "Non", "FK -> disciplines.id", "Discipline pratiquee"],
            ["edition_id", "BIGINT UNSIGNED", "Non", "FK -> editions.id", "Edition de la competition"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    # --- COUCHE CONTROLEUR ---
    doc.add_heading("4. Couche Controleur", level=1)
    doc.add_heading("4.1. Inventaire des Controleurs", level=2)
    add_styled_table(doc,
        ["Controleur", "Methodes", "Description"],
        [
            ["DashboardController", "index()", "Recupere les KPIs, derniers resultats, prochains matchs et meilleurs buteurs pour l'edition selectionnee."],
            ["EditionController", "index, show, create, store, edit, update, destroy", "CRUD complet sur les editions. La methode 'arbre()' genere la vue arborescente du tournoi."],
            ["FaculteController", "index, show, create, store, edit, update, destroy", "CRUD complet sur les facultes."],
            ["DisciplineController", "index, show, create, store, edit, update, destroy", "CRUD complet sur les disciplines sportives."],
            ["EquipeController", "index, show, create, store, edit, update, destroy", "CRUD complet sur les equipes. Inclut la gestion des relations avec Faculte et Discipline."],
            ["JoueurController", "index, create, store, edit, update, destroy", "CRUD complet sur les joueurs. Mappe 'numero' vers 'numero_maillot'."],
            ["MatchController", "index, show, create, store, edit, update, destroy, saisirScore()", "CRUD + saisie de score. La methode saisirScore() gere le JSON des buteurs et l'incrementation des buts."],
            ["ClassementController", "index()", "Calcule dynamiquement le classement par discipline et le top des buteurs."],
            ["UsersManagementController", "index, create, store, edit, update, destroy", "CRUD complet sur les comptes utilisateurs (admin only)."],
            ["ProfileController", "edit, update, destroy", "Gestion du profil personnel de l'utilisateur connecte."],
        ]
    )

    # --- MIDDLEWARES ---
    doc.add_heading("5. Couche Middleware et Securite", level=1)
    doc.add_heading("5.1. Middlewares de Controle d'Acces", level=2)
    add_styled_table(doc,
        ["Middleware", "Cle", "Condition", "Routes Protegees"],
        [
            ["Authentification", "auth", "L'utilisateur doit etre connecte", "Toutes les routes sauf '/' et '/login'"],
            ["Admin Only", "admin", "role === 'admin'", "CRUD Editions, Facultes, Disciplines, Equipes, Gestion Staff"],
            ["Can Manage", "can.manage", "role === 'admin' OR role === 'staff'", "CRUD Joueurs, CRUD Matchs, Saisie de Score"],
            ["CSRF", "VerifyCsrfToken", "Token CSRF valide dans chaque POST/PUT/DELETE", "Tous les formulaires"],
        ]
    )

    doc.add_heading("5.2. Structure de Routage", level=2)
    doc.add_paragraph("Le fichier routes/web.php est organise en trois groupes hierarchiques :")
    doc.add_paragraph("1. Groupe 'admin' : Routes de creation, modification et suppression des entites structurelles (editions, facultes, disciplines, equipes, utilisateurs). Accessible uniquement aux administrateurs.", style='List Number')
    doc.add_paragraph("2. Groupe 'can.manage' : Routes de gestion du contenu operationnel (joueurs, matchs, scores). Accessible aux administrateurs et au staff.", style='List Number')
    doc.add_paragraph("3. Routes publiques (authentifiees) : Routes de consultation en lecture seule (index, show). Accessibles a tout utilisateur connecte.", style='List Number')

    # --- FRONTEND ---
    doc.add_heading("6. Couche Frontend", level=1)
    doc.add_heading("6.1. Systeme de Design", level=2)
    doc.add_paragraph("L'interface utilise un systeme de design unifie defini dans app.css avec les classes CSS suivantes :")
    add_styled_table(doc,
        ["Classe CSS", "Usage"],
        [
            ["enterprise-card", "Conteneur blanc avec bordure, ombre douce et coins arrondis"],
            ["enterprise-btn-primary", "Bouton principal (fond bleu, texte blanc, effet hover)"],
            ["enterprise-btn-secondary", "Bouton secondaire (bordure, fond transparent)"],
            ["enterprise-input", "Champ de saisie standardise (bordure, focus bleu, padding)"],
            ["enterprise-label", "Label de formulaire (majuscules, petit texte, espacement)"],
        ]
    )

    doc.add_heading("6.2. Reactivite avec Alpine.js", level=2)
    doc.add_paragraph("Alpine.js est utilise pour les interactions cote client sans rechargement de page :")
    doc.add_paragraph("Recherche en temps reel : Chaque liste de donnees integre un champ x-model='search' qui filtre dynamiquement les lignes du tableau via x-show.", style='List Bullet')
    doc.add_paragraph("Accordeons : Les panneaux de discipline dans la page des matchs utilisent x-data et @click pour ouvrir/fermer les sections.", style='List Bullet')
    doc.add_paragraph("Ajout dynamique de buteurs : Le formulaire de saisie de score utilise Alpine pour ajouter/retirer des lignes de buteurs sans rechargement.", style='List Bullet')

    doc.save("docs/3_Dossier_Architecture.docx")
    print("[OK] 3_Dossier_Architecture.docx")

# ============================================================
# DOCUMENT 4 : SPECIFICATIONS FONCTIONNELLES
# ============================================================
def create_specifications():
    doc = Document()
    setup_document(doc)
    add_cover_page(doc, "Specifications Fonctionnelles et Techniques", "Regles de Gestion, Cas d'Utilisation, Architecture, Base de Donnees et Securite", "2.0")
    add_toc_placeholder(doc)

    doc.add_heading("1. Regles de Gestion", level=1)
    rules = [
        ("RG01", "Unicite Equipe", "Une faculte ne peut inscrire qu'une seule equipe par discipline et par edition. Le systeme valide cette contrainte a la creation et affiche un message d'erreur en cas de doublon."),
        ("RG02", "Statut d'Edition", "Une edition suit un cycle de vie lineaire : 'a_venir' -> 'en_cours' -> 'terminee'. Lorsqu'une edition est au statut 'terminee', il est impossible de programmer de nouveaux matchs."),
        ("RG03", "Equipes Differentes", "Les deux equipes selectionnees pour un match doivent obligatoirement etre distinctes. La validation 'different:equipe_b_id' empeche la programmation d'un match d'une equipe contre elle-meme."),
        ("RG04", "Calcul des Points", "Le bareme de points est : Victoire = 3 points, Match Nul = 1 point, Defaite = 0 point. Le classement est trie par points decroissants, puis par difference de buts."),
        ("RG05", "Saisie des Buteurs", "Lors de la saisie du score, les buteurs selectionnes doivent appartenir a l'equipe concernee (liste deroulante filtree par equipe). Chaque buteur peut avoir un nombre de buts >= 1."),
        ("RG06", "Mise a Jour des Buts", "A la validation du score, le compteur de buts du joueur dans la table joueurs est automatiquement incremente du nombre de buts saisis pour ce match. Cette operation est irreversible."),
        ("RG07", "Suppression en Cascade", "La suppression d'une equipe entraine la suppression de tous ses joueurs. La suppression d'une edition entraine la suppression de toutes les equipes et matchs associes (dependant de la configuration des cles etrangeres)."),
        ("RG08", "Protection des Editions Terminees", "Les boutons 'Programmer un match' et 'Nouvelle equipe' sont masques pour les editions au statut 'terminee'. Le controleur valide egalement cette regle cote serveur."),
        ("RG09", "Authentification Obligatoire", "L'acces a toute page de la plateforme (hormis la page de connexion et de reinitialisation de mot de passe) necessite une authentification prealable."),
        ("RG10", "Hashage des Mots de Passe", "Tous les mots de passe sont stockes en base de donnees sous forme hashee avec l'algorithme bcrypt. Aucun mot de passe en clair n'est jamais stocke ou affiche."),
    ]
    
    for code, titre, desc in rules:
        doc.add_heading(f"{code} - {titre}", level=3)
        doc.add_paragraph(desc)

    # --- CAS D'UTILISATION ---
    doc.add_heading("2. Cas d'Utilisation Detailles", level=1)

    doc.add_heading("2.1. CU01 - Se Connecter", level=2)
    add_styled_table(doc,
        ["Element", "Description"],
        [
            ["Acteur principal", "Tout utilisateur"],
            ["Precondition", "L'utilisateur possede un compte valide"],
            ["Scenario nominal", "1. L'utilisateur accede a la page de connexion.\n2. Il saisit son email et son mot de passe.\n3. Il clique sur 'Se connecter'.\n4. Le systeme valide les identifiants.\n5. Redirection vers le Tableau de Bord."],
            ["Scenario alternatif", "4a. Les identifiants sont incorrects : message d'erreur 'Ces identifiants ne correspondent pas'."],
            ["Postcondition", "L'utilisateur est authentifie. Une session est creee."],
        ]
    )

    doc.add_heading("2.2. CU02 - Programmer un Match", level=2)
    add_styled_table(doc,
        ["Element", "Description"],
        [
            ["Acteur principal", "Administrateur ou Staff"],
            ["Precondition", "L'utilisateur est connecte. Une edition non terminee existe. Au moins deux equipes sont inscrites."],
            ["Scenario nominal", "1. L'utilisateur clique sur 'Programmer un match'.\n2. Il selectionne l'edition, la discipline, les deux equipes, la date, le lieu et la phase.\n3. Il clique sur 'Enregistrer'.\n4. Le systeme valide les donnees et cree le match avec statut 'planifie'.\n5. Redirection vers la liste des matchs avec message de succes."],
            ["Scenario alternatif", "3a. Les deux equipes selectionnees sont identiques : erreur 'different'.\n3b. L'edition est terminee : erreur 'edition terminee'."],
            ["Postcondition", "Un nouveau match est cree en base de donnees avec le statut 'planifie'."],
        ]
    )

    doc.add_heading("2.3. CU03 - Saisir un Score", level=2)
    add_styled_table(doc,
        ["Element", "Description"],
        [
            ["Acteur principal", "Administrateur ou Staff"],
            ["Precondition", "Le match existe avec le statut 'planifie'. L'utilisateur est connecte avec le role admin ou staff."],
            ["Scenario nominal", "1. L'utilisateur accede a la fiche du match planifie.\n2. Le formulaire de saisie du score s'affiche.\n3. Il saisit le score de chaque equipe.\n4. Il ajoute les buteurs avec le nombre de buts pour chaque joueur.\n5. Il clique sur 'Enregistrer le resultat final'.\n6. Le systeme enregistre le score, stocke les buteurs en JSON, incremente les buts des joueurs et passe le statut a 'joue'."],
            ["Scenario alternatif", "5a. Un joueur inexistant est selectionne : erreur de validation."],
            ["Postcondition", "Le match passe au statut 'joue'. Les classements sont recalcules. Les compteurs de buts des joueurs sont mis a jour."],
        ]
    )

    doc.add_heading("2.4. CU04 - Consulter le Classement", level=2)
    add_styled_table(doc,
        ["Element", "Description"],
        [
            ["Acteur principal", "Tout utilisateur connecte"],
            ["Precondition", "L'utilisateur est connecte. Une edition est selectionnee."],
            ["Scenario nominal", "1. L'utilisateur clique sur 'Classements' dans le menu.\n2. Il selectionne une discipline dans le menu deroulant.\n3. Le tableau de classement s'affiche avec toutes les equipes de cette discipline triees par points.\n4. En dessous, le top 10 des meilleurs buteurs est affiche."],
            ["Postcondition", "Les donnees sont affichees en temps reel (pas de cache)."],
        ]
    )

    doc.add_heading("2.5. CU05 - Gerer les Comptes Utilisateurs", level=2)
    add_styled_table(doc,
        ["Element", "Description"],
        [
            ["Acteur principal", "Administrateur uniquement"],
            ["Precondition", "L'utilisateur est connecte avec le role 'admin'."],
            ["Scenario nominal", "1. L'administrateur accede a 'Gestion Staff'.\n2. Il voit la liste de tous les utilisateurs.\n3. Il clique sur '+ Nouveau Utilisateur'.\n4. Il remplit le formulaire (nom, email, mot de passe, role).\n5. Il clique sur 'Creer le compte'.\n6. Le compte est cree et apparait dans la liste."],
            ["Scenario alternatif", "4a. L'email existe deja : erreur 'L'email a deja ete pris'."],
            ["Postcondition", "Un nouveau compte utilisateur est cree avec le role attribue."],
        ]
    )

    # ================================================================
    # PARTIE II : SPECIFICATIONS TECHNIQUES
    # ================================================================
    doc.add_page_break()
    doc.add_heading("PARTIE II : SPECIFICATIONS TECHNIQUES", level=1)
    doc.add_paragraph("Cette partie presente l'ensemble des choix techniques, l'architecture logicielle, les technologies utilisees, le schema de la base de donnees, la securite et l'environnement de deploiement de la plateforme UniGames.")

    # --- 3. PILE TECHNOLOGIQUE ---
    doc.add_heading("3. Pile Technologique (Technology Stack)", level=1)
    doc.add_heading("3.1. Technologies Backend", level=2)
    add_styled_table(doc,
        ["Technologie", "Version", "Role", "Justification"],
        [
            ["PHP", "8.2+", "Langage de programmation serveur", "Langage mature, performant, avec un ecosysteme riche. Support natif des types, enums et fibres."],
            ["Laravel", "12.x", "Framework applicatif MVC", "Framework PHP le plus populaire. Fournit un ORM (Eloquent), un systeme de routage, des middlewares, la gestion des sessions et un moteur de templates (Blade)."],
            ["Laravel Breeze", "2.x", "Kit d'authentification", "Fournit les fonctionnalites de connexion, deconnexion, inscription, reinitialisation de mot de passe et verification d'email, avec des vues Blade pre-configurees."],
            ["Eloquent ORM", "Inclus dans Laravel", "Mapping Objet-Relationnel", "Permet de manipuler les tables de la base de donnees comme des objets PHP. Gestion automatique des relations (hasMany, belongsTo), des accesseurs calcules et des casts."],
            ["Composer", "2.x", "Gestionnaire de dependances PHP", "Gere l'installation et la mise a jour de toutes les bibliotheques PHP du projet."],
        ]
    )

    doc.add_heading("3.2. Technologies Frontend", level=2)
    add_styled_table(doc,
        ["Technologie", "Version", "Role", "Justification"],
        [
            ["Blade", "Inclus dans Laravel", "Moteur de templates", "Moteur de templates natif de Laravel. Permet l'heritage de layouts, les composants reutilisables (x-app-layout, x-sidebar, x-topbar) et l'injection de donnees."],
            ["Tailwind CSS", "4.x", "Framework CSS utilitaire", "Approche utility-first permettant un design rapide, coherent et responsive. Classes atomiques composables (flex, grid, text-*, bg-*, etc.)."],
            ["Alpine.js", "3.x", "Micro-framework JavaScript reactif", "Fournit la reactivite cote client (x-data, x-model, x-show, @click) sans la complexite d'un framework SPA comme Vue ou React. Ideal pour les interactions legeres."],
            ["Vite", "6.x", "Bundler et serveur de developpement", "Remplacement de Webpack. Compile les assets CSS/JS avec Hot Module Replacement (HMR) pour un rechargement instantane en developpement."],
            ["PostCSS + Autoprefixer", "8.x / 10.x", "Post-traitement CSS", "Assure la compatibilite cross-navigateur en ajoutant automatiquement les prefixes vendeur necessaires."],
        ]
    )

    doc.add_heading("3.3. Base de Donnees", level=2)
    add_styled_table(doc,
        ["Technologie", "Version", "Role", "Justification"],
        [
            ["MySQL", "8.0+", "SGBD Relationnel", "Systeme de gestion de base de donnees relationnelle robuste, performant et largement supporte. Support du JSON natif pour le stockage des buteurs."],
        ]
    )

    doc.add_heading("3.4. Outils de Developpement", level=2)
    add_styled_table(doc,
        ["Outil", "Role"],
        [
            ["Git", "Controle de version distribue (depot GitHub)"],
            ["GitHub", "Hebergement du depot distant et collaboration"],
            ["Node.js / NPM", "Environnement d'execution JavaScript pour le build frontend"],
            ["PHP Artisan", "Interface en ligne de commande (CLI) de Laravel pour les migrations, seeders, generation de controleurs, etc."],
            ["VS Code", "Editeur de code principal avec extensions PHP, Blade et Tailwind CSS"],
        ]
    )

    # --- 4. ARCHITECTURE MVC ---
    doc.add_heading("4. Architecture Logicielle (MVC)", level=1)
    doc.add_heading("4.1. Patron Model-View-Controller", level=2)
    doc.add_paragraph("UniGames adopte strictement le patron MVC impose par Laravel, qui separe les responsabilites en trois couches distinctes :")
    add_bold_para(doc, "Modele (Model) : ", "Represente les entites metier et encapsule la logique d'acces aux donnees via Eloquent ORM. Chaque modele correspond a une table de la base de donnees et definit les relations, accesseurs calcules et regles de validation.")
    add_bold_para(doc, "Vue (View) : ", "Les templates Blade (.blade.php) generent le HTML. Ils utilisent le systeme de composants Laravel (x-app-layout, x-sidebar, x-topbar) pour la coherence visuelle. Le CSS est gere par Tailwind, et les interactions legeres par Alpine.js.")
    add_bold_para(doc, "Controleur (Controller) : ", "Orchestre le flux applicatif : reception des requetes HTTP, validation des donnees d'entree, appel aux modeles, preparation des donnees et retour de la vue appropriee.")

    doc.add_heading("4.2. Diagramme de Flux d'une Requete HTTP", level=2)
    doc.add_paragraph("Le cycle de vie d'une requete dans UniGames suit le flux suivant :")
    doc.add_paragraph("1. Le navigateur envoie une requete HTTP (GET/POST/PUT/DELETE) au serveur.", style='List Number')
    doc.add_paragraph("2. Le routeur (routes/web.php) identifie la route correspondante et le controleur associe.", style='List Number')
    doc.add_paragraph("3. Les middlewares sont executes dans l'ordre (authentification, verification du role, protection CSRF).", style='List Number')
    doc.add_paragraph("4. Le controleur recoit la requete validee, interagit avec le(s) modele(s) Eloquent.", style='List Number')
    doc.add_paragraph("5. Le modele Eloquent execute les requetes SQL via le Query Builder et retourne les resultats.", style='List Number')
    doc.add_paragraph("6. Le controleur transmet les donnees a la vue Blade.", style='List Number')
    doc.add_paragraph("7. Blade compile le template en HTML, intègre les assets CSS/JS via Vite, et renvoie la reponse au navigateur.", style='List Number')

    # --- 5. MODELES ET BASE DE DONNEES ---
    doc.add_heading("5. Schema de la Base de Donnees", level=1)
    doc.add_heading("5.1. Modele Conceptuel de Donnees (MCD)", level=2)
    doc.add_paragraph("Le systeme repose sur 7 entites metier principales interconnectees par des relations de type 1:N (un-a-plusieurs) :")
    add_styled_table(doc,
        ["Entite", "Table SQL", "Cle Primaire", "Description"],
        [
            ["Utilisateur", "users", "id (BIGINT UNSIGNED, auto-increment)", "Compte utilisateur avec role (admin ou staff)"],
            ["Edition", "editions", "id (BIGINT UNSIGNED, auto-increment)", "Competition / tournoi avec dates et statut"],
            ["Faculte", "facultes", "id (BIGINT UNSIGNED, auto-increment)", "Etablissement universitaire participant"],
            ["Discipline", "disciplines", "id (BIGINT UNSIGNED, auto-increment)", "Sport pratique dans le tournoi"],
            ["Equipe", "equipes", "id (BIGINT UNSIGNED, auto-increment)", "Equipe rattachee a une faculte, discipline et edition"],
            ["Joueur", "joueurs", "id (BIGINT UNSIGNED, auto-increment)", "Sportif rattache a une equipe"],
            ["Match", "matchs", "id (BIGINT UNSIGNED, auto-increment)", "Rencontre entre deux equipes avec score et buteurs"],
        ]
    )

    doc.add_heading("5.2. Relations entre Entites", level=2)
    add_styled_table(doc,
        ["Relation", "Type", "Cle Etrangere", "Contrainte"],
        [
            ["Edition -> Faculte", "1:N (Une edition a plusieurs facultes)", "facultes.edition_id -> editions.id", "CASCADE on DELETE"],
            ["Edition -> Equipe", "1:N (Une edition a plusieurs equipes)", "equipes.edition_id -> editions.id", "CASCADE on DELETE"],
            ["Edition -> Match", "1:N (Une edition a plusieurs matchs)", "matchs.edition_id -> editions.id", "CASCADE on DELETE"],
            ["Faculte -> Equipe", "1:N (Une faculte a plusieurs equipes)", "equipes.faculte_id -> facultes.id", "CASCADE on DELETE"],
            ["Discipline -> Equipe", "1:N (Une discipline a plusieurs equipes)", "equipes.discipline_id -> disciplines.id", "CASCADE on DELETE"],
            ["Discipline -> Match", "1:N (Une discipline a plusieurs matchs)", "matchs.discipline_id -> disciplines.id", "CASCADE on DELETE"],
            ["Equipe -> Joueur", "1:N (Une equipe a plusieurs joueurs)", "joueurs.equipe_id -> equipes.id", "CASCADE on DELETE"],
            ["Equipe -> Match (A)", "1:N (Une equipe joue plusieurs matchs en tant qu'equipe A)", "matchs.equipe_a_id -> equipes.id", "CASCADE on DELETE"],
            ["Equipe -> Match (B)", "1:N (Une equipe joue plusieurs matchs en tant qu'equipe B)", "matchs.equipe_b_id -> equipes.id", "CASCADE on DELETE"],
        ]
    )

    doc.add_heading("5.3. Dictionnaire de Donnees Detaille", level=2)

    doc.add_paragraph("Table : users (Utilisateurs)")
    add_styled_table(doc,
        ["Colonne", "Type SQL", "Nullable", "Contrainte", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK, Auto-increment", "Identifiant unique de l'utilisateur"],
            ["name", "VARCHAR(255)", "Non", "-", "Nom complet de l'utilisateur"],
            ["email", "VARCHAR(255)", "Non", "UNIQUE", "Adresse e-mail (identifiant de connexion)"],
            ["email_verified_at", "TIMESTAMP", "Oui", "-", "Date de verification de l'email"],
            ["password", "VARCHAR(255)", "Non", "-", "Mot de passe hashe (bcrypt, 12 rounds)"],
            ["role", "VARCHAR(50)", "Non", "DEFAULT 'staff'", "Role : 'admin' ou 'staff'"],
            ["remember_token", "VARCHAR(100)", "Oui", "-", "Token 'Se souvenir de moi'"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Date de creation du compte"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Date de derniere modification"],
        ]
    )

    doc.add_paragraph("Table : editions (Editions du Tournoi)")
    add_styled_table(doc,
        ["Colonne", "Type SQL", "Nullable", "Contrainte", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK, Auto-increment", "Identifiant unique"],
            ["nom", "VARCHAR(255)", "Non", "-", "Nom de l'edition (ex: 'Edition 2026')"],
            ["date_debut", "DATE", "Non", "-", "Date de debut du tournoi"],
            ["date_fin", "DATE", "Non", "-", "Date de fin du tournoi"],
            ["lieu", "VARCHAR(255)", "Oui", "-", "Lieu principal du tournoi"],
            ["description", "TEXT", "Oui", "-", "Description libre"],
            ["statut", "VARCHAR(50)", "Non", "DEFAULT 'a_venir'", "Cycle de vie : a_venir -> en_cours -> terminee"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    doc.add_paragraph("Table : facultes (Facultes / Etablissements)")
    add_styled_table(doc,
        ["Colonne", "Type SQL", "Nullable", "Contrainte", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK, Auto-increment", "Identifiant unique"],
            ["nom", "VARCHAR(255)", "Non", "-", "Nom complet de la faculte"],
            ["logo", "VARCHAR(255)", "Oui", "-", "Chemin vers le fichier logo"],
            ["couleur", "VARCHAR(7)", "Oui", "-", "Code couleur hexadecimal (#RRGGBB)"],
            ["edition_id", "BIGINT UNSIGNED", "Non", "FK -> editions.id", "Edition de rattachement"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    doc.add_paragraph("Table : disciplines (Disciplines Sportives)")
    add_styled_table(doc,
        ["Colonne", "Type SQL", "Nullable", "Contrainte", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK, Auto-increment", "Identifiant unique"],
            ["nom", "VARCHAR(255)", "Non", "-", "Nom du sport (ex: Football, Basketball)"],
            ["type", "VARCHAR(255)", "Oui", "-", "Type de sport (collectif, individuel)"],
            ["nb_joueurs", "INT", "Oui", "-", "Nombre de joueurs par equipe"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    doc.add_paragraph("Table : equipes (Equipes)")
    add_styled_table(doc,
        ["Colonne", "Type SQL", "Nullable", "Contrainte", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK, Auto-increment", "Identifiant unique"],
            ["nom", "VARCHAR(255)", "Non", "-", "Nom de l'equipe"],
            ["faculte_id", "BIGINT UNSIGNED", "Non", "FK -> facultes.id", "Faculte representee"],
            ["discipline_id", "BIGINT UNSIGNED", "Non", "FK -> disciplines.id", "Discipline pratiquee"],
            ["edition_id", "BIGINT UNSIGNED", "Non", "FK -> editions.id", "Edition de la competition"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    doc.add_paragraph("Table : joueurs (Joueurs)")
    add_styled_table(doc,
        ["Colonne", "Type SQL", "Nullable", "Contrainte", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK, Auto-increment", "Identifiant unique"],
            ["nom", "VARCHAR(255)", "Non", "-", "Nom de famille du joueur"],
            ["prenom", "VARCHAR(255)", "Non", "-", "Prenom du joueur"],
            ["sexe", "CHAR(1)", "Non", "-", "Genre : M (Masculin) ou F (Feminin)"],
            ["equipe_id", "BIGINT UNSIGNED", "Non", "FK -> equipes.id", "Equipe d'appartenance"],
            ["numero_maillot", "INT", "Oui", "-", "Numero de maillot du joueur"],
            ["buts", "INT", "Non", "DEFAULT 0", "Compteur total de buts marques (incremente automatiquement)"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    doc.add_paragraph("Table : matchs (Rencontres Sportives)")
    add_styled_table(doc,
        ["Colonne", "Type SQL", "Nullable", "Contrainte", "Description"],
        [
            ["id", "BIGINT UNSIGNED", "Non", "PK, Auto-increment", "Identifiant unique"],
            ["equipe_a_id", "BIGINT UNSIGNED", "Non", "FK -> equipes.id", "Equipe recevante (domicile)"],
            ["equipe_b_id", "BIGINT UNSIGNED", "Non", "FK -> equipes.id", "Equipe visiteuse (exterieur)"],
            ["discipline_id", "BIGINT UNSIGNED", "Non", "FK -> disciplines.id", "Discipline sportive du match"],
            ["edition_id", "BIGINT UNSIGNED", "Non", "FK -> editions.id", "Edition du tournoi"],
            ["date_match", "DATETIME", "Non", "-", "Date et heure prevue du match"],
            ["lieu", "VARCHAR(255)", "Oui", "-", "Lieu / stade du match"],
            ["phase", "VARCHAR(255)", "Non", "-", "Phase : Poules, Quarts, Demies, Petite Finale, Finale"],
            ["score_a", "INT", "Oui", "DEFAULT NULL", "Score de l'equipe A (null si non joue)"],
            ["score_b", "INT", "Oui", "DEFAULT NULL", "Score de l'equipe B (null si non joue)"],
            ["statut", "VARCHAR(50)", "Non", "DEFAULT 'planifie'", "Cycle : planifie -> en_cours -> joue"],
            ["buteurs", "JSON", "Oui", "-", "Donnees JSON des buteurs ({equipe_a: [...], equipe_b: [...]})"],
            ["created_at", "TIMESTAMP", "Oui", "-", "Horodatage creation"],
            ["updated_at", "TIMESTAMP", "Oui", "-", "Horodatage modification"],
        ]
    )

    # --- 6. ROUTES ET API ---
    doc.add_heading("6. Cartographie des Routes (API Interne)", level=1)
    doc.add_heading("6.1. Routes Publiques (Authentifiees)", level=2)
    doc.add_paragraph("Ces routes sont accessibles a tout utilisateur connecte (admin ou staff) :")
    add_styled_table(doc,
        ["Methode HTTP", "URI", "Controleur@Methode", "Description"],
        [
            ["GET", "/dashboard", "DashboardController@index", "Tableau de bord avec KPIs"],
            ["GET", "/editions", "EditionController@index", "Liste des editions"],
            ["GET", "/editions/{id}", "EditionController@show", "Detail d'une edition"],
            ["GET", "/editions/{id}/arbre", "EditionController@arbre", "Arbre du tournoi"],
            ["GET", "/facultes", "FaculteController@index", "Liste des facultes"],
            ["GET", "/facultes/{id}", "FaculteController@show", "Detail d'une faculte"],
            ["GET", "/disciplines", "DisciplineController@index", "Liste des disciplines"],
            ["GET", "/equipes", "EquipeController@index", "Liste des equipes"],
            ["GET", "/equipes/{id}", "EquipeController@show", "Detail d'une equipe"],
            ["GET", "/joueurs", "JoueurController@index", "Liste des joueurs"],
            ["GET", "/matchs", "MatchController@index", "Calendrier des matchs"],
            ["GET", "/matchs/{id}", "MatchController@show", "Detail d'un match"],
            ["GET", "/classements", "ClassementController@index", "Classements par discipline"],
        ]
    )

    doc.add_heading("6.2. Routes Protegees (Admin uniquement)", level=2)
    doc.add_paragraph("Ces routes necessitent le role 'admin' et sont protegees par le middleware 'admin' :")
    add_styled_table(doc,
        ["Methode HTTP", "URI", "Action", "Description"],
        [
            ["GET", "/editions/create", "CRUD", "Formulaire de creation d'edition"],
            ["POST", "/editions", "CRUD", "Enregistrer une nouvelle edition"],
            ["GET", "/editions/{id}/edit", "CRUD", "Formulaire de modification d'edition"],
            ["PUT", "/editions/{id}", "CRUD", "Mettre a jour une edition"],
            ["DELETE", "/editions/{id}", "CRUD", "Supprimer une edition"],
            ["GET/POST", "/facultes/create, /facultes", "CRUD", "Creation de faculte"],
            ["PUT/DELETE", "/facultes/{id}", "CRUD", "Modification/Suppression de faculte"],
            ["GET/POST", "/disciplines/create, /disciplines", "CRUD", "Creation de discipline"],
            ["PUT/DELETE", "/disciplines/{id}", "CRUD", "Modification/Suppression de discipline"],
            ["GET/POST", "/equipes/create, /equipes", "CRUD", "Creation d'equipe"],
            ["PUT/DELETE", "/equipes/{id}", "CRUD", "Modification/Suppression d'equipe"],
            ["GET/POST", "/users/create, /users", "CRUD", "Creation de compte utilisateur"],
            ["PUT/DELETE", "/users/{id}", "CRUD", "Modification/Suppression de compte"],
        ]
    )

    doc.add_heading("6.3. Routes Protegees (Admin + Staff)", level=2)
    doc.add_paragraph("Ces routes necessitent le role 'admin' ou 'staff' et sont protegees par le middleware 'can.manage' :")
    add_styled_table(doc,
        ["Methode HTTP", "URI", "Action", "Description"],
        [
            ["GET/POST", "/joueurs/create, /joueurs", "CRUD", "Creation de joueur"],
            ["GET/PUT", "/joueurs/{id}/edit, /joueurs/{id}", "CRUD", "Modification de joueur"],
            ["DELETE", "/joueurs/{id}", "CRUD", "Suppression de joueur"],
            ["GET/POST", "/matchs/create, /matchs", "CRUD", "Programmation de match"],
            ["GET/PUT", "/matchs/{id}/edit, /matchs/{id}", "CRUD", "Modification de match"],
            ["DELETE", "/matchs/{id}", "CRUD", "Suppression de match"],
            ["POST", "/matchs/{id}/score", "Specifique", "Saisie du score et des buteurs d'un match"],
        ]
    )

    # --- 7. SECURITE ---
    doc.add_heading("7. Securite et Controle d'Acces", level=1)
    doc.add_heading("7.1. Authentification", level=2)
    doc.add_paragraph("Le systeme d'authentification est base sur Laravel Breeze et offre les fonctionnalites suivantes :")
    doc.add_paragraph("Connexion par email et mot de passe avec protection contre les attaques par force brute (throttling).", style='List Bullet')
    doc.add_paragraph("Sessions stockees en base de donnees (table 'sessions') avec duree de vie de 120 minutes.", style='List Bullet')
    doc.add_paragraph("Token 'Se souvenir de moi' (remember_token) pour les sessions persistantes.", style='List Bullet')
    doc.add_paragraph("Reinitialisation de mot de passe par email avec token a usage unique.", style='List Bullet')
    doc.add_paragraph("Deconnexion avec invalidation complete de la session.", style='List Bullet')

    doc.add_heading("7.2. Hashage des Mots de Passe", level=2)
    doc.add_paragraph("Tous les mots de passe sont hashes avec l'algorithme bcrypt configure a 12 rounds (BCRYPT_ROUNDS=12 dans .env). Aucun mot de passe en clair n'est jamais stocke, affiche ou transmis en reponse HTTP.")

    doc.add_heading("7.3. Protection CSRF", level=2)
    doc.add_paragraph("Chaque formulaire HTML inclut un token CSRF via la directive Blade @csrf. Le middleware VerifyCsrfToken valide ce token pour toutes les requetes POST, PUT, PATCH et DELETE. En cas d'echec, une erreur HTTP 419 'Page Expired' est retournee.")

    doc.add_heading("7.4. Matrice des Roles et Permissions", level=2)
    add_styled_table(doc,
        ["Fonctionnalite", "Admin", "Staff", "Non-authentifie"],
        [
            ["Voir le tableau de bord", "Oui", "Oui", "Non (redirige vers login)"],
            ["Consulter editions, facultes, equipes, joueurs, matchs", "Oui", "Oui", "Non"],
            ["Consulter les classements et l'arbre du tournoi", "Oui", "Oui", "Non"],
            ["Creer/Modifier/Supprimer des editions", "Oui", "Non (403 Forbidden)", "Non"],
            ["Creer/Modifier/Supprimer des facultes", "Oui", "Non", "Non"],
            ["Creer/Modifier/Supprimer des disciplines", "Oui", "Non", "Non"],
            ["Creer/Modifier/Supprimer des equipes", "Oui", "Non", "Non"],
            ["Creer/Modifier/Supprimer des joueurs", "Oui", "Oui", "Non"],
            ["Programmer/Modifier/Supprimer des matchs", "Oui", "Oui", "Non"],
            ["Saisir le score d'un match", "Oui", "Oui", "Non"],
            ["Gerer les comptes utilisateurs", "Oui", "Non", "Non"],
            ["Modifier son propre profil", "Oui", "Oui", "Non"],
        ]
    )

    doc.add_heading("7.5. Validation des Donnees", level=2)
    doc.add_paragraph("Toutes les donnees envoyees par les formulaires sont validees cote serveur par les controleurs Laravel via la methode $request->validate(). Les regles de validation incluent :")
    add_styled_table(doc,
        ["Champ", "Regles de Validation", "Message d'Erreur"],
        [
            ["email", "required, email, unique:users", "L'email a deja ete pris / L'email est obligatoire"],
            ["password", "required, min:8, confirmed", "Le mot de passe doit contenir au moins 8 caracteres"],
            ["nom (edition/faculte)", "required, string, max:255", "Le nom est obligatoire"],
            ["date_debut", "required, date", "La date de debut est obligatoire"],
            ["equipe_a_id", "required, exists:equipes,id, different:equipe_b_id", "Les deux equipes doivent etre differentes"],
            ["score_a / score_b", "required, integer, min:0", "Le score doit etre un nombre positif ou nul"],
        ]
    )

    # --- 8. FRONTEND TECHNIQUE ---
    doc.add_heading("8. Specifications Techniques Frontend", level=1)
    doc.add_heading("8.1. Composants Blade Reutilisables", level=2)
    add_styled_table(doc,
        ["Composant", "Fichier", "Role"],
        [
            ["x-app-layout", "layouts/app.blade.php", "Layout principal avec sidebar, topbar et zone de contenu"],
            ["x-sidebar", "components/sidebar.blade.php", "Barre laterale de navigation avec liens actifs dynamiques"],
            ["x-topbar", "components/topbar.blade.php", "Barre superieure avec titre de page, actions et avatar utilisateur"],
            ["x-input-error", "components/input-error.blade.php", "Affichage des erreurs de validation sous les champs de formulaire"],
        ]
    )

    doc.add_heading("8.2. Systeme de Classes CSS Enterprise", level=2)
    doc.add_paragraph("Le design system UniGames est base sur un ensemble de classes CSS personnalisees definies dans resources/css/app.css :")
    add_styled_table(doc,
        ["Classe CSS", "Proprietes Principales", "Usage"],
        [
            ["enterprise-card", "bg: white, border-radius: 12px, border: 1px solid slate-200, box-shadow: 0 1px 3px", "Conteneur generique pour cartes, formulaires, tableaux"],
            ["enterprise-btn-primary", "bg: indigo-600, color: white, padding: 10px 20px, border-radius: 8px, hover: indigo-700", "Boutons d'action principale (Creer, Enregistrer)"],
            ["enterprise-btn-secondary", "bg: transparent, border: slate-200, color: slate-700, hover: slate-50", "Boutons secondaires (Annuler, Retour)"],
            ["enterprise-input", "border: slate-200, bg: slate-50, focus: border-indigo, height: 42px, transition: 150ms", "Champs de saisie de formulaire"],
            ["enterprise-label", "text-transform: uppercase, font-size: 11px, font-weight: bold, letter-spacing: widest", "Labels de formulaire"],
        ]
    )

    doc.add_heading("8.3. Interactions Alpine.js", level=2)
    doc.add_paragraph("Alpine.js est utilise pour 3 types d'interactions cote client :")
    doc.add_paragraph("Recherche en temps reel : Les listes de donnees (joueurs, equipes, matchs) integrent un champ x-model='search' lie a un filtre x-show qui masque dynamiquement les lignes ne correspondant pas a la saisie.", style='List Bullet')
    doc.add_paragraph("Accordeons : Les panneaux de discipline sur la page des matchs utilisent x-data='{open: false}' et @click='open = !open' pour replier/deplier les sections.", style='List Bullet')
    doc.add_paragraph("Formulaire dynamique de buteurs : Le formulaire de saisie de score utilise x-data avec un tableau reactif de buteurs. Les fonctions addButeur() et removeButeur() permettent d'ajouter/retirer des lignes sans rechargement de page.", style='List Bullet')

    # --- 9. ENVIRONNEMENT DE DEPLOIEMENT ---
    doc.add_heading("9. Environnement de Deploiement", level=1)
    doc.add_heading("9.1. Configuration Serveur Requise", level=2)
    add_styled_table(doc,
        ["Composant", "Specification Minimale", "Recommandation"],
        [
            ["Systeme d'exploitation", "Windows 10 / Ubuntu 20.04 / macOS 12+", "Ubuntu 22.04 LTS en production"],
            ["PHP", "8.2+", "8.3 pour les performances optimales"],
            ["MySQL", "8.0+", "8.0.33+ avec InnoDB"],
            ["Node.js", "18+", "20 LTS"],
            ["RAM", "2 Go minimum", "4 Go recommande"],
            ["Disque", "500 Mo pour l'application", "SSD recommande pour la base de donnees"],
        ]
    )

    doc.add_heading("9.2. Variables d'Environnement Critiques", level=2)
    add_styled_table(doc,
        ["Variable", "Valeur Dev", "Valeur Production", "Description"],
        [
            ["APP_ENV", "local", "production", "Environnement d'execution"],
            ["APP_DEBUG", "true", "false", "Affichage des erreurs detaillees (JAMAIS true en production)"],
            ["APP_KEY", "base64:...", "base64:...", "Cle de chiffrement (generee par php artisan key:generate)"],
            ["APP_URL", "http://localhost:8001", "https://votre-domaine.com", "URL de base de l'application"],
            ["DB_CONNECTION", "mysql", "mysql", "Pilote de base de donnees"],
            ["DB_HOST", "127.0.0.1", "adresse-serveur-db", "Adresse du serveur MySQL"],
            ["DB_DATABASE", "unigames", "unigames", "Nom de la base de donnees"],
            ["BCRYPT_ROUNDS", "12", "12", "Nombre de rounds de hachage bcrypt"],
            ["SESSION_DRIVER", "database", "database", "Stockage des sessions"],
            ["SESSION_LIFETIME", "120", "120", "Duree de vie des sessions en minutes"],
        ]
    )

    doc.add_heading("9.3. Commandes de Deploiement", level=2)
    doc.add_paragraph("Procedure de mise en production :")
    doc.add_paragraph("1. git clone https://github.com/Johnny-jonnes/Unigames_Laravel.git", style='List Number')
    doc.add_paragraph("2. cd Unigames_Laravel && composer install --optimize-autoloader --no-dev", style='List Number')
    doc.add_paragraph("3. npm install && npm run build", style='List Number')
    doc.add_paragraph("4. cp .env.example .env && php artisan key:generate", style='List Number')
    doc.add_paragraph("5. Configurer les variables DB_* dans le fichier .env", style='List Number')
    doc.add_paragraph("6. php artisan migrate --force", style='List Number')
    doc.add_paragraph("7. php artisan config:cache && php artisan route:cache && php artisan view:cache", style='List Number')
    doc.add_paragraph("8. Configurer le serveur web (Apache/Nginx) pour pointer vers le dossier /public", style='List Number')

    doc.save("docs/4_Specifications_Fonctionnelles_et_Techniques.docx")
    print("[OK] 4_Specifications_Fonctionnelles_et_Techniques.docx")

# ============================================================
# DOCUMENT 5 : GUIDE D'INSTALLATION
# ============================================================
def create_installation():
    doc = Document()
    setup_document(doc)
    add_cover_page(doc, "Guide de Deploiement", "Procedure Complete d'Installation et de Configuration", "2.0")
    add_toc_placeholder(doc)

    doc.add_heading("1. Prerequis Systeme", level=1)
    add_styled_table(doc,
        ["Logiciel", "Version Minimale", "Role", "Lien de Telechargement"],
        [
            ["PHP", "8.2+", "Moteur d'execution du code Laravel", "https://www.php.net/downloads"],
            ["Composer", "2.x", "Gestionnaire de dependances PHP", "https://getcomposer.org/download/"],
            ["MySQL", "8.0+", "Systeme de gestion de base de donnees relationnelle", "https://dev.mysql.com/downloads/"],
            ["Node.js", "18+", "Environnement d'execution JavaScript pour le build frontend", "https://nodejs.org/"],
            ["NPM", "9+", "Gestionnaire de paquets Node.js (inclus avec Node.js)", "Inclus avec Node.js"],
            ["Git", "2.x", "Systeme de controle de version", "https://git-scm.com/downloads"],
        ]
    )

    doc.add_heading("1.1. Extensions PHP Requises", level=2)
    exts = ["pdo_mysql", "mbstring", "tokenizer", "xml", "ctype", "json", "bcmath", "openssl", "fileinfo"]
    for ext in exts:
        doc.add_paragraph(f"php-{ext}", style='List Bullet')

    doc.add_heading("2. Procedure d'Installation Pas a Pas", level=1)

    doc.add_heading("2.1. Etape 1 : Cloner le Depot Git", level=2)
    doc.add_paragraph("Ouvrir un terminal et executer :")
    doc.add_paragraph("git clone https://github.com/Johnny-jonnes/Unigames_Laravel.git")
    doc.add_paragraph("cd Unigames_Laravel")

    doc.add_heading("2.2. Etape 2 : Installer les Dependances PHP", level=2)
    doc.add_paragraph("composer install")
    doc.add_paragraph("Cette commande telecharge et installe toutes les bibliotheques PHP definies dans composer.json (Laravel, Breeze, etc.).")

    doc.add_heading("2.3. Etape 3 : Installer les Dependances Frontend", level=2)
    doc.add_paragraph("npm install")
    doc.add_paragraph("Cette commande installe les paquets JavaScript necessaires (Vite, Tailwind CSS, Alpine.js, PostCSS, Autoprefixer).")

    doc.add_heading("2.4. Etape 4 : Configurer l'Environnement", level=2)
    doc.add_paragraph("Copier le fichier de configuration :")
    doc.add_paragraph("cp .env.example .env    (Linux/Mac)")
    doc.add_paragraph("copy .env.example .env  (Windows)")
    doc.add_paragraph("")
    doc.add_paragraph("Modifier les parametres suivants dans le fichier .env :")
    add_styled_table(doc,
        ["Variable", "Valeur", "Description"],
        [
            ["APP_NAME", "UniGames", "Nom de l'application affiche dans la barre de titre"],
            ["APP_URL", "http://localhost:8001", "URL de base de l'application"],
            ["DB_CONNECTION", "mysql", "Type de base de donnees"],
            ["DB_HOST", "127.0.0.1", "Adresse du serveur MySQL"],
            ["DB_PORT", "3306", "Port MySQL"],
            ["DB_DATABASE", "unigames", "Nom de la base de donnees (a creer manuellement)"],
            ["DB_USERNAME", "root", "Nom d'utilisateur MySQL"],
            ["DB_PASSWORD", "(votre mot de passe)", "Mot de passe de l'utilisateur MySQL"],
        ]
    )

    doc.add_heading("2.5. Etape 5 : Creer la Base de Donnees", level=2)
    doc.add_paragraph("Se connecter a MySQL et creer la base :")
    doc.add_paragraph("mysql -u root -p")
    doc.add_paragraph("CREATE DATABASE unigames CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;")
    doc.add_paragraph("EXIT;")

    doc.add_heading("2.6. Etape 6 : Generer la Cle d'Application", level=2)
    doc.add_paragraph("php artisan key:generate")
    doc.add_paragraph("Cette commande genere une cle de chiffrement unique dans le fichier .env (APP_KEY). Elle est indispensable pour le chiffrement des cookies et des sessions.")

    doc.add_heading("2.7. Etape 7 : Executer les Migrations", level=2)
    doc.add_paragraph("php artisan migrate")
    doc.add_paragraph("Cette commande cree toutes les tables dans la base de donnees selon les fichiers de migration situes dans database/migrations/.")

    doc.add_heading("2.8. Etape 8 : Peupler la Base de Donnees (Optionnel)", level=2)
    doc.add_paragraph("Pour injecter des donnees de test realistes (universites guineennes, joueurs, matchs avec scores) :")
    doc.add_paragraph("php artisan db:seed --class=GuineanDataSeeder")
    doc.add_paragraph("php artisan db:seed --class=FixEdition2026Seeder")
    doc.add_paragraph("Ces seeders creent une edition 2026 complete avec des facultes guineennes, des equipes, des joueurs aux noms realistes et des matchs avec scores simules.")

    doc.add_heading("2.9. Etape 9 : Lancer les Serveurs de Developpement", level=2)
    doc.add_paragraph("Ouvrir deux terminaux separes :")
    doc.add_paragraph("Terminal 1 (Serveur PHP) : php artisan serve --port=8001")
    doc.add_paragraph("Terminal 2 (Serveur Vite) : npm run dev")
    doc.add_paragraph("Acceder a l'application : http://localhost:8001")

    doc.add_heading("3. Comptes par Defaut", level=1)
    add_styled_table(doc,
        ["Role", "Email", "Mot de Passe"],
        [
            ["Administrateur", "admin@unigames.com", "password"],
            ["Staff", "(cree par l'admin)", "(defini par l'admin)"],
        ]
    )

    doc.add_heading("4. Depannage Courant", level=1)
    add_styled_table(doc,
        ["Probleme", "Solution"],
        [
            ["Page blanche apres connexion", "Verifier que le serveur Vite tourne (npm run dev). Sans Vite, les assets CSS/JS ne sont pas compiles."],
            ["Erreur SQLSTATE", "Verifier les parametres DB_* dans .env. S'assurer que la base 'unigames' existe et que MySQL est demarre."],
            ["Erreur 419 Page Expired", "Erreur CSRF. Vider le cache : php artisan cache:clear et recharger la page."],
            ["Erreur 403 Forbidden", "L'utilisateur n'a pas les droits suffisants. Verifier le role dans la table users."],
            ["Erreur 500 Column not found", "Executer les migrations : php artisan migrate. Verifier que toutes les colonnes existent."],
        ]
    )

    doc.save("docs/5_Guide_Installation.docx")
    print("[OK] 5_Guide_Installation.docx")

# ============================================================
# DOCUMENT 6 : PLAN DE TESTS
# ============================================================
def create_plan_test():
    doc = Document()
    setup_document(doc)
    add_cover_page(doc, "Plan de Tests", "Cahier de Recette, Scenarios de Validation et Resultats", "2.0")
    add_toc_placeholder(doc)

    doc.add_heading("1. Strategie de Test", level=1)
    doc.add_heading("1.1. Approche", level=2)
    doc.add_paragraph("La strategie de test d'UniGames repose sur deux niveaux complementaires :")
    doc.add_paragraph("Tests fonctionnels manuels : Validation des parcours utilisateurs critiques via un cahier de recette detaille.", style='List Bullet')
    doc.add_paragraph("Tests automatises (unitaires et d'integration) : Fournis par Laravel Breeze pour les fonctionnalites d'authentification (tests/Feature/Auth/).", style='List Bullet')

    doc.add_heading("1.2. Environnement de Test", level=2)
    add_styled_table(doc,
        ["Element", "Valeur"],
        [
            ["Navigateur principal", "Google Chrome (derniere version)"],
            ["Navigateurs secondaires", "Firefox, Edge"],
            ["Systeme d'exploitation", "Windows 11"],
            ["Base de donnees", "MySQL 8.0 (base 'unigames')"],
            ["Donnees de test", "Seeders GuineanDataSeeder + FixEdition2026Seeder"],
        ]
    )

    doc.add_heading("2. Cahier de Recette Detaille", level=1)

    tests = [
        ("CT-AUTH-01", "Connexion avec identifiants valides", "Administrateur",
         "1. Acceder a http://localhost:8001\n2. Saisir admin@unigames.com / password\n3. Cliquer sur 'Se connecter'",
         "Redirection vers le Tableau de Bord. Le nom 'Admin' apparait dans la topbar.", "VALIDE"),
        
        ("CT-AUTH-02", "Connexion avec identifiants invalides", "Tout utilisateur",
         "1. Acceder a la page de connexion\n2. Saisir un email ou mot de passe incorrect\n3. Cliquer sur 'Se connecter'",
         "Message d'erreur rouge : 'Ces identifiants ne correspondent pas a nos enregistrements.'", "VALIDE"),
        
        ("CT-AUTH-03", "Deconnexion", "Tout utilisateur connecte",
         "1. Cliquer sur l'avatar en haut a droite\n2. Cliquer sur 'Deconnexion'",
         "Redirection vers la page de connexion. La session est detruite.", "VALIDE"),
        
        ("CT-DASH-01", "Affichage du Tableau de Bord avec edition", "Tout utilisateur",
         "1. Se connecter\n2. Selectionner 'Edition 2026' dans le selecteur",
         "Les KPIs s'affichent (Facultes, Equipes, Matchs, Joueurs). Les derniers resultats et prochains matchs sont visibles.", "VALIDE"),
        
        ("CT-EDIT-01", "Creer une nouvelle edition", "Administrateur",
         "1. Aller dans Editions\n2. Cliquer '+ Nouvelle Edition'\n3. Remplir : Nom='Edition 2027', Dates, Statut='a_venir'\n4. Soumettre",
         "L'edition apparait dans la liste avec le badge 'A venir'.", "VALIDE"),
        
        ("CT-EDIT-02", "Modifier une edition", "Administrateur",
         "1. Depuis la liste des editions, cliquer 'Modifier'\n2. Changer le statut en 'en_cours'\n3. Soumettre",
         "Le statut de l'edition est mis a jour. Le badge change de couleur.", "VALIDE"),
        
        ("CT-FAC-01", "Creer une faculte", "Administrateur",
         "1. Aller dans Facultes\n2. Cliquer '+ Nouvelle Faculte'\n3. Remplir le nom et la couleur\n4. Soumettre",
         "La faculte apparait dans la liste avec les initiales comme logo.", "VALIDE"),
        
        ("CT-EQU-01", "Inscrire une equipe", "Administrateur",
         "1. Aller dans Equipes\n2. Cliquer '+ Nouvelle Equipe'\n3. Selectionner la faculte, la discipline et l'edition\n4. Soumettre",
         "L'equipe est creee et associee a la bonne faculte/discipline/edition.", "VALIDE"),
        
        ("CT-JOU-01", "Ajouter un joueur", "Staff ou Administrateur",
         "1. Aller dans Joueurs\n2. Cliquer '+ Nouveau Joueur'\n3. Remplir nom, prenom, sexe, equipe, numero\n4. Soumettre",
         "Le joueur apparait dans la liste avec toutes ses informations.", "VALIDE"),
        
        ("CT-JOU-02", "Modifier un joueur", "Staff ou Administrateur",
         "1. Dans la liste des joueurs, cliquer 'Modifier'\n2. Changer le numero de maillot\n3. Soumettre",
         "Le joueur est mis a jour avec le nouveau numero.", "VALIDE"),
        
        ("CT-JOU-03", "Supprimer un joueur", "Staff ou Administrateur",
         "1. Dans la liste des joueurs, cliquer le bouton rouge 'Supprimer'\n2. Confirmer dans la boite de dialogue",
         "Le joueur disparait de la liste. Message de succes affiche.", "VALIDE"),
        
        ("CT-MAT-01", "Programmer un match", "Staff ou Administrateur",
         "1. Aller dans Matchs\n2. Cliquer '+ Programmer un match'\n3. Remplir tous les champs obligatoires\n4. Soumettre",
         "Le match apparait dans la discipline correspondante avec le badge 'Planifie'.", "VALIDE"),
        
        ("CT-MAT-02", "Saisir un score avec buteurs", "Staff ou Administrateur",
         "1. Ouvrir un match 'planifie'\n2. Saisir score_a=2, score_b=1\n3. Ajouter 2 buteurs pour l'equipe A\n4. Ajouter 1 buteur pour l'equipe B\n5. Soumettre",
         "Le match passe au statut 'Joue'. Le score est affiche. Les buteurs sont listes. Les compteurs de buts des joueurs sont incrementes.", "VALIDE"),
        
        ("CT-MAT-03", "Supprimer un match", "Staff ou Administrateur",
         "1. Dans la liste des matchs, cliquer le bouton rouge 'Supprimer'\n2. Confirmer la suppression",
         "Le match disparait de la liste. Redirection vers la liste avec message de succes.", "VALIDE"),
        
        ("CT-CLA-01", "Consulter le classement", "Tout utilisateur",
         "1. Aller dans Classements\n2. Selectionner la discipline 'Football'",
         "Le tableau de classement s'affiche avec colonnes MJ, V, N, D, BM, BE, DB, Pts. Le tri est correct.", "VALIDE"),
        
        ("CT-CLA-02", "Meilleurs buteurs", "Tout utilisateur",
         "1. Aller dans Classements\n2. Consulter la section 'Meilleurs Buteurs'",
         "Le top 10 des buteurs est affiche avec nom, equipe, faculte et nombre de buts.", "VALIDE"),
        
        ("CT-ARB-01", "Arbre du tournoi", "Tout utilisateur",
         "1. Aller dans Editions\n2. Cliquer sur l'edition 2026\n3. Cliquer 'Arbre du Tournoi'",
         "L'arbre s'affiche avec les phases (Poules, Quarts, Demies, Petite Finale, Grande Finale). Les scores sont visibles. Le champion a une couronne.", "VALIDE"),
        
        ("CT-SEC-01", "Acces non autorise (Staff vers Admin)", "Staff",
         "1. Se connecter en tant que staff\n2. Acceder directement a /editions/create par URL",
         "Redirection ou page d'erreur 403 Forbidden.", "VALIDE"),
        
        ("CT-SEC-02", "Acces non authentifie", "Anonyme",
         "1. Sans etre connecte, acceder directement a /dashboard",
         "Redirection automatique vers la page de connexion.", "VALIDE"),
        
        ("CT-USR-01", "Creer un compte staff", "Administrateur",
         "1. Aller dans Gestion Staff\n2. Cliquer '+ Nouveau Utilisateur'\n3. Remplir nom, email, mot de passe, role='staff'\n4. Soumettre",
         "Le compte apparait dans la liste avec le badge 'Staff'.", "VALIDE"),
        
        ("CT-PRF-01", "Modifier son profil", "Tout utilisateur",
         "1. Cliquer sur l'avatar\n2. Aller dans 'Profil'\n3. Modifier le nom\n4. Soumettre",
         "Le nom est mis a jour. Message de succes affiche.", "VALIDE"),
    ]

    for code, titre, acteur, actions, resultat, statut in tests:
        doc.add_heading(f"{code} : {titre}", level=3)
        add_styled_table(doc,
            ["Element", "Detail"],
            [
                ["Code", code],
                ["Titre", titre],
                ["Acteur", acteur],
                ["Actions", actions],
                ["Resultat Attendu", resultat],
                ["Statut", statut],
            ]
        )

    # --- SYNTHESE ---
    doc.add_heading("3. Synthese des Resultats", level=1)
    add_styled_table(doc,
        ["Categorie", "Nb Tests", "Valides", "Echoues", "Taux de Reussite"],
        [
            ["Authentification", "3", "3", "0", "100%"],
            ["Tableau de Bord", "1", "1", "0", "100%"],
            ["Editions", "2", "2", "0", "100%"],
            ["Facultes", "1", "1", "0", "100%"],
            ["Equipes", "1", "1", "0", "100%"],
            ["Joueurs", "3", "3", "0", "100%"],
            ["Matchs", "3", "3", "0", "100%"],
            ["Classements", "2", "2", "0", "100%"],
            ["Arbre du Tournoi", "1", "1", "0", "100%"],
            ["Securite", "2", "2", "0", "100%"],
            ["Gestion Utilisateurs", "1", "1", "0", "100%"],
            ["Profil", "1", "1", "0", "100%"],
            ["TOTAL", "21", "21", "0", "100%"],
        ]
    )

    doc.save("docs/6_Plan_Tests.docx")
    print("[OK] 6_Plan_Tests.docx")

# ============================================================
# DOCUMENT 7 : DOSSIER DE MAQUETTAGE
# ============================================================
def create_maquette():
    doc = Document()
    setup_document(doc)
    add_cover_page(doc, "Dossier de Maquettage", "Charte Graphique, Design System et Description des Interfaces", "2.0")
    add_toc_placeholder(doc)

    doc.add_heading("1. Charte Graphique", level=1)
    doc.add_heading("1.1. Palette de Couleurs", level=2)
    add_styled_table(doc,
        ["Nom", "Code Hex", "Variable CSS", "Usage"],
        [
            ["Bleu Principal", "#4F46E5 (Indigo 600)", "--color-primary", "Boutons principaux, liens actifs, sidebar selectionnee"],
            ["Vert Accent", "#10B981 (Emerald 500)", "--color-accent", "Badges de succes, scores positifs, indicateurs"],
            ["Rouge Danger", "#EF4444 (Red 500)", "--color-danger", "Boutons de suppression, erreurs de validation"],
            ["Texte Principal", "#1E293B (Slate 800)", "--text-primary", "Titres, texte courant"],
            ["Texte Secondaire", "#64748B (Slate 500)", "--text-muted", "Labels, informations complementaires"],
            ["Fond de Page", "#F8FAFC (Slate 50)", "--bg-page", "Arriere-plan general de l'application"],
            ["Fond de Carte", "#FFFFFF", "--bg-card", "Conteneurs, formulaires, tableaux"],
            ["Bordure", "#E2E8F0 (Slate 200)", "--border-color", "Bordures de cartes, separateurs"],
        ]
    )

    doc.add_heading("1.2. Typographie", level=2)
    add_styled_table(doc,
        ["Element", "Police", "Taille", "Graisse", "Couleur"],
        [
            ["Titre de Page (h1)", "Calibri/Inter", "20px", "Bold (700)", "Slate 800"],
            ["Sous-titre (h2)", "Calibri/Inter", "14-16px", "Bold (700)", "Slate 800"],
            ["Label de Section", "Calibri/Inter", "11px", "Bold (700) + Uppercase", "Slate 500"],
            ["Texte Courant", "Calibri/Inter", "13px", "Normal (400)", "Slate 700"],
            ["Badge/Etiquette", "Calibri/Inter", "10-11px", "Bold (700)", "Variable selon type"],
            ["Donnees Numeriques", "Mono (font-mono)", "14-36px", "Bold (700-900)", "Variable"],
        ]
    )

    doc.add_heading("1.3. Iconographie", level=2)
    doc.add_paragraph("Les icones sont des SVG inline (sans dependance externe). Chaque icone utilise les attributs stroke='currentColor' et fill='none' pour s'adapter automatiquement a la couleur du texte parent. Les tailles standard sont w-4 h-4 (petite), w-5 h-5 (normale), et w-16 h-16 (illustration vide).")

    # --- COMPOSANTS ---
    doc.add_heading("2. Systeme de Design (Composants UI)", level=1)
    doc.add_heading("2.1. Carte Enterprise (enterprise-card)", level=2)
    doc.add_paragraph("Composant de base pour tous les conteneurs de contenu. Caracteristiques :")
    doc.add_paragraph("Fond blanc avec coins arrondis (border-radius: 12px)", style='List Bullet')
    doc.add_paragraph("Bordure fine de couleur Slate 200", style='List Bullet')
    doc.add_paragraph("Ombre subtile (box-shadow: 0 1px 3px rgba(0,0,0,0.1))", style='List Bullet')
    doc.add_paragraph("Transition douce au survol (changement d'ombre et de bordure)", style='List Bullet')

    doc.add_heading("2.2. Boutons", level=2)
    add_styled_table(doc,
        ["Classe", "Apparence", "Usage"],
        [
            ["enterprise-btn-primary", "Fond Indigo 600, texte blanc, coins arrondis, ombre douce. Au survol : fond Indigo 700.", "Actions principales (Creer, Enregistrer, Valider)"],
            ["enterprise-btn-secondary", "Fond transparent, bordure Slate 200, texte Slate 700. Au survol : fond Slate 50.", "Actions secondaires (Annuler, Retour, Modifier)"],
            ["Bouton Danger (inline)", "Fond transparent, bordure Rouge 200, texte Rouge 600. Au survol : fond Rouge 50.", "Actions de suppression"],
        ]
    )

    doc.add_heading("2.3. Champs de Formulaire", level=2)
    doc.add_paragraph("Classe enterprise-input : Champ avec bordure Slate 200, fond Slate 50 au repos, bordure Indigo au focus. Hauteur standardisee de 42px. Taille de texte 13px. Transition douce (150ms) sur la couleur de bordure.")
    doc.add_paragraph("Classe enterprise-label : Label en majuscules, taille 11px, graisse Bold, couleur Slate 500, avec espacement de lettres augmente (tracking-widest).")

    # --- DESCRIPTION DES PAGES ---
    doc.add_heading("3. Description Detaillee des Interfaces", level=1)

    doc.add_heading("3.1. Page de Connexion", level=2)
    doc.add_paragraph("Layout : Centrage vertical et horizontal sur fond degrade bleu-indigo avec motifs de fond subtils.")
    doc.add_paragraph("Contenu : Carte blanche centree contenant le logo UniGames, le titre 'Connexion', deux champs de saisie (email et mot de passe), une case a cocher 'Se souvenir de moi', un bouton primaire 'Se connecter' pleine largeur, et un lien 'Mot de passe oublie ?' en dessous.")
    doc.add_paragraph("Experience utilisateur : Le champ email recoit le focus automatiquement a l'ouverture. Les erreurs de validation s'affichent en rouge sous chaque champ concerne.")
    if os.path.exists("docs/images/2_login.jpeg"):
        doc.add_picture("docs/images/2_login.jpeg", width=Inches(6.0))

    doc.add_heading("3.2. Tableau de Bord", level=2)
    doc.add_paragraph("Section haute : Selecteur d'edition en haut a droite (dropdown). Grille de 4 a 8 cartes KPI alignees horizontalement, chacune avec une icone, un chiffre en gras et un libelle descriptif.")
    doc.add_paragraph("Section milieu : Deux colonnes - 'Derniers Resultats' (5 matchs les plus recents avec scores) et 'Prochains Matchs' (5 matchs planifies avec dates). Chaque ligne est cliquable.")
    doc.add_paragraph("Section basse : Tableau 'Meilleurs Buteurs' avec classement des 5 premiers joueurs. Les donnees incluent le nom, l'equipe, la faculte et le nombre de buts.")
    if os.path.exists("docs/images/4_dashboard_principal.jpeg"):
        doc.add_picture("docs/images/4_dashboard_principal.jpeg", width=Inches(6.0))

    doc.add_heading("3.3. Page des Matchs", level=2)
    doc.add_paragraph("Structure en accordeons : Chaque discipline sportive possede un panneau repliable. Le panneau ouvert affiche un tableau complet de tous les matchs.")
    doc.add_paragraph("Barre de recherche : En haut, un champ de recherche en temps reel filtre les matchs par nom d'equipe ou lieu.")
    doc.add_paragraph("Tableau des matchs : Colonnes Date/Lieu, Equipe A, Score, Equipe B, Statut (badge colore), Actions (Voir/Modifier/Supprimer).")
    doc.add_paragraph("Ligne cliquable : Chaque ligne du tableau est cliquable et redirige vers la fiche detaillee du match. Les boutons d'action utilisent event.stopPropagation() pour eviter les conflits de clic.")
    if os.path.exists("docs/images/13_calendrier_rencontres.jpeg"):
        doc.add_picture("docs/images/13_calendrier_rencontres.jpeg", width=Inches(6.0))

    doc.add_heading("3.4. Fiche Match (Score)", level=2)
    doc.add_paragraph("Scoreboard central : Grande carte avec les logos/initiales des deux equipes, le score en format XXL (36px monospace) et un badge 'Termine' (vert). Les buteurs de chaque equipe sont listes sous le score.")
    doc.add_paragraph("Informations contextuelles : Grille 3 colonnes affichant Date, Lieu et Phase.")
    doc.add_paragraph("Formulaire de saisie (si match planifie) : Deux colonnes avec champ numerique geant pour le score et section dynamique d'ajout de buteurs via Alpine.js.")
    if os.path.exists("docs/images/14_details_match.jpeg"):
        doc.add_picture("docs/images/14_details_match.jpeg", width=Inches(6.0))

    doc.add_heading("3.5. Arbre du Tournoi", level=2)
    doc.add_paragraph("Disposition horizontale avec defilement : Les phases sont disposees de gauche a droite (Poules -> Quarts -> Demies -> Petite Finale -> Grande Finale). Chaque phase est une colonne verticale de cartes de matchs.")
    doc.add_paragraph("Cartes de match : Style different selon la phase. Les matchs de poules sont compacts. Les demi-finales ont un fond colore pour le vainqueur. La grande finale est mise en valeur avec un degrade ambre, une ombre prononcee et un badge 'OR'.")
    if os.path.exists("docs/images/18_arbre_tournoi.jpeg"):
        doc.add_picture("docs/images/18_arbre_tournoi.jpeg", width=Inches(6.0))

    doc.add_heading("3.6. Page des Classements", level=2)
    doc.add_paragraph("Selecteur de discipline : Menu deroulant en haut permettant de choisir la discipline sportive.")
    doc.add_paragraph("Tableau de classement : Style zebra (lignes alternees), en-tete bleue. Colonnes : #, Equipe (avec faculte), MJ, V, N, D, BM, BE, DB, Pts. La premiere ligne est mise en evidence.")
    doc.add_paragraph("Meilleurs buteurs : Tableau secondaire en dessous avec le top 10, incluant la photo/initiale du joueur, son nom, equipe, faculte et nombre de buts.")
    if os.path.exists("docs/images/15_classements_statistiques.jpeg"):
        doc.add_picture("docs/images/15_classements_statistiques.jpeg", width=Inches(6.0))

    doc.add_heading("4. Composants de Navigation", level=1)
    doc.add_heading("4.1. Sidebar (Barre Laterale)", level=2)
    doc.add_paragraph("Position fixe a gauche, largeur de 250px, fond sombre (Slate 900). Le logo UniGames est affiche en haut. Les liens de navigation sont organises verticalement avec icones SVG et labels textuels. Le lien actif est mis en evidence par un fond semi-transparent et un indicateur colore a gauche.")
    
    doc.add_heading("4.2. Topbar (Barre Superieure)", level=2)
    doc.add_paragraph("Barre horizontale en haut de la zone de contenu. Elle contient le titre de la page courante a gauche et les actions contextuelles a droite (boutons de creation, selecteur d'edition). L'avatar de l'utilisateur connecte est affiche a l'extreme droite avec un menu deroulant (Profil, Deconnexion).")

    doc.save("docs/7_Dossier_Maquettage.docx")
    print("[OK] 7_Dossier_Maquettage.docx")

def add_screenshots(doc, title="Captures d'Ecran"):
    import glob
    images = glob.glob("docs/images/*.jpeg")
    if not images:
        images = glob.glob("docs/images/*.png")
    
    if images:
        doc.add_heading(title, level=1)
        doc.add_paragraph("Les images ci-dessous presentent l'interface reelle de l'application.")
        for img in sorted(images):
            try:
                doc.add_picture(img, width=Inches(6.0))
                doc.add_paragraph() # espace apres l'image
            except Exception as e:
                print(f"Erreur lors de l'ajout de l'image {img}: {e}")

# ============================================================
# MAIN
# ============================================================
def main():
    if not os.path.exists('docs'):
        os.makedirs('docs')
    
    print("=" * 60)
    print("  GENERATION DES DOCUMENTS PROFESSIONNELS UniGames")
    print("=" * 60)
    
    create_cahier_des_charges()
    create_manuel_utilisateur()
    create_dossier_architecture()
    create_specifications()
    create_installation()
    create_plan_test()
    create_maquette()
    
    print("=" * 60)
    print("  7 documents generes avec succes dans le dossier 'docs/'")
    print("=" * 60)

if __name__ == "__main__":
    main()
